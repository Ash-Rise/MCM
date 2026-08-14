from __future__ import annotations

import argparse
import hashlib
import heapq
import json
import math
import re
from collections import Counter, deque
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Callable, Iterable

import numpy as np
import pandas as pd
from docx import Document
from scipy.integrate import quad
from scipy.optimize import Bounds, LinearConstraint, linprog, milp


MINUTES_PER_DAY = 1440.0
BUSY_MINUTES = 45.0
PREP_MINUTES = 3.0
SPEED_KMH = 45.0
DAILY_CAP = 12
DAILY_CALLS = 140
GOLDEN_RESPONSE_MINUTES = 4.0
SERVICE_RADIUS_KM = SPEED_KMH * GOLDEN_RESPONSE_MINUTES / 60.0
DELAY_PENALTY_YUAN_PER_MINUTE = 200.0
EPS = 1e-9
STATEMENT_FILENAME = "problem-a-ambulance-dispatch-statement.docx"


def solution_root() -> Path:
    return Path(__file__).resolve().parents[1]


def problem_statement_path(project_root: Path | None = None) -> Path:
    root = (project_root or solution_root()).resolve()
    return root.parents[1] / "problem-statements" / STATEMENT_FILENAME


@dataclass(frozen=True)
class ProblemData:
    zone_ids: tuple[int, ...]
    zone_names: tuple[str, ...]
    zone_xy: np.ndarray
    area: np.ndarray
    population: np.ndarray
    demand: np.ndarray
    hospital_distance: np.ndarray
    site_ids: tuple[str, ...]
    site_xy: np.ndarray
    site_caps: np.ndarray
    distance: np.ndarray

    @property
    def demand_density(self) -> np.ndarray:
        return self.demand / self.area


@dataclass(frozen=True)
class Call:
    call_id: int
    arrival_min: float
    zone: int


@dataclass
class Ambulance:
    ambulance_id: int
    site: int
    local_id: int
    busy_until: float = 0.0
    day_count: int = 0
    reserve: bool = False
    activation_min: float = 0.0
    external: bool = False


@dataclass(frozen=True)
class DispatchRecord:
    call_id: int
    zone: int
    ambulance_id: int
    site: int
    arrival_min: float
    dispatch_min: float
    response_min: float
    wait_min: float
    dispatch_day: int
    c_loss_min: float | None


def _number(text: str) -> float:
    match = re.search(r"[-+]?\d+(?:\.\d+)?", text.replace(",", ""))
    if not match:
        raise ValueError(f"No numeric value in cell: {text!r}")
    return float(match.group())


def read_problem(docx_path: Path) -> ProblemData:
    document = Document(docx_path)
    if len(document.tables) < 3:
        raise ValueError("A-problem input must contain three tables")

    zone_rows = document.tables[0].rows[1:]
    site_rows = document.tables[1].rows[1:]
    if len(zone_rows) != 10 or len(site_rows) != 6:
        raise ValueError("Expected 10 demand zones and 6 candidate sites")

    zone_ids = tuple(int(_number(row.cells[0].text)) for row in zone_rows)
    zone_names = tuple(row.cells[1].text.strip() for row in zone_rows)
    zone_xy = np.array(
        [[_number(row.cells[2].text), _number(row.cells[3].text)] for row in zone_rows],
        dtype=float,
    )
    area = np.array([_number(row.cells[4].text) for row in zone_rows], dtype=float)
    population = np.array([_number(row.cells[5].text) for row in zone_rows], dtype=float)
    demand = np.array([_number(row.cells[6].text) for row in zone_rows], dtype=float)
    hospital_distance = np.array([_number(row.cells[7].text) for row in zone_rows], dtype=float)
    site_ids = tuple(row.cells[0].text.strip() for row in site_rows)
    site_xy = np.array(
        [[_number(row.cells[1].text), _number(row.cells[2].text)] for row in site_rows],
        dtype=float,
    )
    site_caps = np.array([int(_number(row.cells[3].text)) for row in site_rows], dtype=int)
    distance = np.linalg.norm(zone_xy[:, None, :] - site_xy[None, :, :], axis=2)

    if not math.isclose(float(demand.sum()), 140.0):
        raise ValueError("Daily demand does not sum to 140")
    if int(site_caps.sum()) != 12:
        raise ValueError("Site capacities do not sum to 12")
    return ProblemData(
        zone_ids=zone_ids,
        zone_names=zone_names,
        zone_xy=zone_xy,
        area=area,
        population=population,
        demand=demand,
        hospital_distance=hospital_distance,
        site_ids=site_ids,
        site_xy=site_xy,
        site_caps=site_caps,
        distance=distance,
    )


def solve_q1(data: ProblemData) -> dict[str, object]:
    n_zones, n_sites = data.distance.shape
    n_x = n_zones * n_sites
    n_vars = n_x + 2 * n_sites
    v_offset = n_x
    y_offset = n_x + n_sites

    objective = np.zeros(n_vars)
    objective[:n_x] = data.distance.ravel()
    rows: list[np.ndarray] = []
    lower: list[float] = []
    upper: list[float] = []

    for i in range(n_zones):
        row = np.zeros(n_vars)
        row[i * n_sites : (i + 1) * n_sites] = 1.0
        rows.append(row)
        lower.append(data.demand[i])
        upper.append(data.demand[i])
    for j in range(n_sites):
        row = np.zeros(n_vars)
        row[j:n_x:n_sites] = 1.0
        row[v_offset + j] = -DAILY_CAP
        rows.append(row)
        lower.append(-np.inf)
        upper.append(0.0)

        row = np.zeros(n_vars)
        row[y_offset + j] = 1.0
        row[v_offset + j] = -1.0
        rows.append(row)
        lower.append(-np.inf)
        upper.append(0.0)

        row = np.zeros(n_vars)
        row[v_offset + j] = 1.0
        row[y_offset + j] = -data.site_caps[j]
        rows.append(row)
        lower.append(-np.inf)
        upper.append(0.0)

    row = np.zeros(n_vars)
    row[v_offset : v_offset + n_sites] = 1.0
    rows.append(row)
    lower.append(-np.inf)
    upper.append(12.0)

    lower_bounds = np.zeros(n_vars)
    upper_bounds = np.full(n_vars, np.inf)
    upper_bounds[v_offset : v_offset + n_sites] = data.site_caps
    upper_bounds[y_offset:] = 1.0
    integrality = np.zeros(n_vars, dtype=int)
    integrality[v_offset:] = 1
    mip = milp(
        c=objective,
        integrality=integrality,
        bounds=Bounds(lower_bounds, upper_bounds),
        constraints=LinearConstraint(np.vstack(rows), np.array(lower), np.array(upper)),
        options={"time_limit": 30.0},
    )
    if not mip.success:
        raise RuntimeError(f"Q1 MILP failed: {mip.message}")
    vehicles = np.rint(mip.x[v_offset : v_offset + n_sites]).astype(int)
    opened = np.rint(mip.x[y_offset:]).astype(int)

    a_eq = np.zeros((n_zones, n_x))
    for i in range(n_zones):
        a_eq[i, i * n_sites : (i + 1) * n_sites] = 1.0
    a_ub = np.zeros((n_sites, n_x))
    for j in range(n_sites):
        a_ub[j, j:n_x:n_sites] = 1.0
    site_capacity = DAILY_CAP * vehicles
    phase1 = linprog(
        data.distance.ravel(),
        A_ub=a_ub,
        b_ub=site_capacity,
        A_eq=a_eq,
        b_eq=data.demand,
        bounds=(0, None),
        method="highs",
    )
    if not phase1.success:
        raise RuntimeError(f"Q1 transport LP failed: {phase1.message}")
    d_star = float(phase1.fun)

    service = (data.distance <= SERVICE_RADIUS_KM + EPS).astype(float).ravel()
    phase2 = linprog(
        -service,
        A_ub=np.vstack([a_ub, data.distance.ravel()]),
        b_ub=np.r_[site_capacity, d_star + 1e-8],
        A_eq=a_eq,
        b_eq=data.demand,
        bounds=(0, None),
        method="highs",
    )
    if not phase2.success:
        raise RuntimeError(f"Q1 lexicographic LP failed: {phase2.message}")
    assignment = phase2.x.reshape(n_zones, n_sites)
    rounded = np.rint(assignment)
    assignment[np.abs(assignment - rounded) < 1e-6] = rounded[np.abs(assignment - rounded) < 1e-6]
    assignment[np.abs(assignment) < 1e-7] = 0.0
    loads = assignment.sum(axis=0)
    rounded_distance = float(np.sum(assignment * data.distance))
    if rounded_distance > d_star + 1e-7:
        raise AssertionError("Cleaning the transport solution changed its primary objective")
    service_covered = float((assignment * (data.distance <= SERVICE_RADIUS_KM + EPS)).sum())
    strict_center_proxy = float((assignment * (data.distance <= 0.75 + EPS)).sum())
    potential_3km = float(
        data.demand[(data.distance.min(axis=1) <= SERVICE_RADIUS_KM + EPS)].sum()
    )

    return {
        "milp_status": int(mip.status),
        "phase1_status": int(phase1.status),
        "phase2_status": int(phase2.status),
        "vehicles": vehicles,
        "opened": opened,
        "assignment": assignment,
        "loads": loads,
        "distance_total": d_star,
        "distance_mean": d_star / data.demand.sum(),
        "static_response_mean": PREP_MINUTES + (60.0 / SPEED_KMH) * d_star / data.demand.sum(),
        "service_radius_km": SERVICE_RADIUS_KM,
        "service_3km_coverage": service_covered / data.demand.sum(),
        "strict_center_proxy_coverage": strict_center_proxy / data.demand.sum(),
        "potential_3km_coverage": potential_3km / data.demand.sum(),
        "max_demand_residual": float(np.max(np.abs(assignment.sum(axis=1) - data.demand))),
        "max_capacity_violation": float(np.max(np.maximum(loads - site_capacity, 0.0))),
        "mip_lp_objective_gap": abs(float(mip.fun) - d_star),
    }


def _periodic_gaussian(hour: float | np.ndarray, mu: float, sigma: float) -> float | np.ndarray:
    if np.isscalar(hour):
        scalar_hour = float(hour)
        return sum(
            math.exp(-((scalar_hour - mu + 24.0 * shift) ** 2) / (2.0 * sigma**2))
            for shift in range(-3, 4)
        )

    hour_array = np.asarray(hour, dtype=float)
    value = np.zeros_like(hour_array)
    for shift in range(-3, 4):
        value += np.exp(-((hour_array - mu + 24.0 * shift) ** 2) / (2.0 * sigma**2))
    return value


def _raw_intraday(hour: float | np.ndarray) -> float | np.ndarray:
    return (
        1.0
        + 0.8 * _periodic_gaussian(hour, 9.0, 2.0)
        + 1.0 * _periodic_gaussian(hour, 18.0, 2.5)
    )


INTRADAY_NORM = quad(lambda x: float(_raw_intraday(x)), 0.0, 24.0, epsabs=1e-12)[0]


def intraday_density(hour: float | np.ndarray) -> float | np.ndarray:
    density = _raw_intraday(np.mod(hour, 24.0)) / INTRADAY_NORM
    if np.isscalar(hour):
        return float(density)
    return density


def delay_penalty_cost(response_minutes: float | np.ndarray) -> float | np.ndarray:
    response = np.asarray(response_minutes, dtype=float)
    cost = DELAY_PENALTY_YUAN_PER_MINUTE * np.maximum(
        response - GOLDEN_RESPONSE_MINUTES,
        0.0,
    )
    if np.isscalar(response_minutes):
        return float(cost)
    return cost


def generate_calls(data: ProblemData, days: int, seed: int) -> list[Call]:
    if days <= 0:
        raise ValueError("Simulation days must be positive")
    rng = np.random.default_rng(seed)
    grid = np.linspace(0.0, 24.0, 24 * 60 + 1)
    density_upper = 1.001 * float(np.max(intraday_density(grid)))
    arrivals: list[tuple[float, int]] = []
    zone_probability = data.demand / data.demand.sum()
    for day in range(days):
        accepted_hours: list[float] = []
        while len(accepted_hours) < DAILY_CALLS:
            remaining = DAILY_CALLS - len(accepted_hours)
            candidates = rng.uniform(0.0, 24.0, size=max(remaining * 2, 32))
            acceptance = rng.uniform(0.0, density_upper, size=len(candidates))
            accepted = candidates[acceptance <= intraday_density(candidates)]
            accepted_hours.extend(accepted[:remaining].tolist())
        for hour in sorted(accepted_hours):
            zone = int(rng.choice(len(data.zone_ids), p=zone_probability))
            arrivals.append((MINUTES_PER_DAY * day + 60.0 * hour, zone))
    return [Call(i, arrival, zone) for i, (arrival, zone) in enumerate(arrivals)]


def build_fleet(
    data: ProblemData,
    reserve_vector: Iterable[int] | None = None,
    external_sites: Iterable[int] | None = None,
    external_activation_min: float = 0.0,
) -> list[Ambulance]:
    reserve = list(reserve_vector or [0] * len(data.site_ids))
    if len(reserve) != len(data.site_ids):
        raise ValueError("Reserve vector length does not match site count")
    if not math.isfinite(external_activation_min) or external_activation_min < 0.0:
        raise ValueError("External activation time must be finite and nonnegative")
    fleet: list[Ambulance] = []
    ambulance_id = 0
    for site, count in enumerate(data.site_caps):
        if reserve[site] < 0 or reserve[site] > count - 1:
            raise ValueError(f"Invalid reserve count at site {site}")
        for local_id in range(int(count)):
            fleet.append(
                Ambulance(
                    ambulance_id=ambulance_id,
                    site=site,
                    local_id=local_id,
                    reserve=local_id < reserve[site],
                )
            )
            ambulance_id += 1
    external_local_ids = [int(count) for count in data.site_caps]
    for raw_site in external_sites or ():
        site = int(raw_site)
        if site != raw_site or not 0 <= site < len(data.site_ids):
            raise ValueError(f"Invalid external ambulance site: {raw_site}")
        fleet.append(
            Ambulance(
                ambulance_id=ambulance_id,
                site=site,
                local_id=external_local_ids[site],
                activation_min=float(external_activation_min),
                external=True,
            )
        )
        external_local_ids[site] += 1
        ambulance_id += 1
    return fleet


def _current_candidates(fleet: list[Ambulance], time_min: float) -> list[Ambulance]:
    return [
        ambulance
        for ambulance in fleet
        if ambulance.activation_min <= time_min + EPS
        and ambulance.busy_until <= time_min + EPS
        and ambulance.day_count < DAILY_CAP
    ]


def _next_midnight(time_min: float) -> float:
    return (math.floor(time_min / MINUTES_PER_DAY) + 1) * MINUTES_PER_DAY


def _known_wait(
    ambulance: Ambulance,
    future_min: float,
    state_time_min: float,
    busy_override: float | None = None,
    count_override: int | None = None,
) -> float:
    if ambulance.activation_min > state_time_min + EPS:
        return math.inf
    busy_until = ambulance.busy_until if busy_override is None else busy_override
    day_count = ambulance.day_count if count_override is None else count_override
    state_day = int(math.floor((state_time_min + EPS) / MINUTES_PER_DAY))
    future_day = int(math.floor((future_min + EPS) / MINUTES_PER_DAY))
    if future_day == state_day and day_count >= DAILY_CAP:
        eligible = max(_next_midnight(state_time_min), busy_until, ambulance.activation_min)
    else:
        eligible = max(future_min, busy_until, ambulance.activation_min)
    return max(0.0, eligible - future_min)


def _predicted_zone_response(
    data: ProblemData,
    fleet: list[Ambulance],
    zone: int,
    future_min: float,
    state_time_min: float,
    dispatched: Ambulance | None,
) -> float:
    best = math.inf
    for ambulance in fleet:
        busy_override = None
        count_override = None
        if dispatched is not None and ambulance.ambulance_id == dispatched.ambulance_id:
            busy_override = state_time_min + BUSY_MINUTES
            count_override = ambulance.day_count + 1
        wait = _known_wait(
            ambulance,
            future_min,
            state_time_min,
            busy_override=busy_override,
            count_override=count_override,
        )
        response = wait + PREP_MINUTES + 60.0 * data.distance[zone, ambulance.site] / SPEED_KMH
        best = min(best, response)
    return best


_GAUSS_NODES, _GAUSS_WEIGHTS = np.polynomial.legendre.leggauss(8)


def cumulative_response_loss(
    data: ProblemData,
    fleet: list[Ambulance],
    dispatched: Ambulance,
    time_min: float,
    rate_multiplier: Callable[[float], np.ndarray] | None = None,
) -> float:
    start = time_min
    end = time_min + BUSY_MINUTES
    total = 0.0
    breakpoints = {start, end}
    midnight = _next_midnight(start)
    if start + EPS < midnight < end - EPS:
        breakpoints.add(midnight)
    for ambulance in fleet:
        if start + EPS < ambulance.busy_until < end - EPS:
            breakpoints.add(ambulance.busy_until)
        if start + EPS < ambulance.activation_min < end - EPS:
            breakpoints.add(ambulance.activation_min)
    ordered = sorted(breakpoints)
    for left, right in zip(ordered, ordered[1:]):
        nodes = 0.5 * (right - left) * _GAUSS_NODES + 0.5 * (left + right)
        weights_in_hours = 0.5 * (right - left) * _GAUSS_WEIGHTS / 60.0
        for future_min, weight in zip(nodes, weights_in_hours, strict=True):
            multiplier = 1.0 if rate_multiplier is None else rate_multiplier(float(future_min))
            rates = data.demand * float(intraday_density(future_min / 60.0)) * multiplier
            response_increase = np.empty(len(data.zone_ids))
            for zone in range(len(data.zone_ids)):
                baseline = _predicted_zone_response(data, fleet, zone, future_min, time_min, None)
                after = _predicted_zone_response(data, fleet, zone, future_min, time_min, dispatched)
                response_increase[zone] = max(0.0, after - baseline)
            total += weight * float(np.dot(rates, response_increase))
    return max(0.0, total)


def cumulative_response_losses(
    data: ProblemData,
    fleet: list[Ambulance],
    candidates: Iterable[Ambulance],
    time_min: float,
    rate_multiplier: Callable[[float], np.ndarray] | None = None,
) -> dict[int, float]:
    candidate_list = list(candidates)
    if not candidate_list:
        return {}
    fleet_index = {ambulance.ambulance_id: index for index, ambulance in enumerate(fleet)}
    totals = {ambulance.ambulance_id: 0.0 for ambulance in candidate_list}
    start = time_min
    end = time_min + BUSY_MINUTES
    breakpoints = {start, end}
    midnight = _next_midnight(start)
    if start + EPS < midnight < end - EPS:
        breakpoints.add(midnight)
    for ambulance in fleet:
        if start + EPS < ambulance.busy_until < end - EPS:
            breakpoints.add(ambulance.busy_until)
        if start + EPS < ambulance.activation_min < end - EPS:
            breakpoints.add(ambulance.activation_min)

    ordered = sorted(breakpoints)
    candidate_indices = np.array([fleet_index[a.ambulance_id] for a in candidate_list], dtype=int)
    busy_until = np.array([ambulance.busy_until for ambulance in fleet], dtype=float)
    day_count = np.array([ambulance.day_count for ambulance in fleet], dtype=int)
    activation_min = np.array([ambulance.activation_min for ambulance in fleet], dtype=float)
    travel = PREP_MINUTES + 60.0 * data.distance[:, [ambulance.site for ambulance in fleet]].T / SPEED_KMH
    for left, right in zip(ordered, ordered[1:]):
        nodes = 0.5 * (right - left) * _GAUSS_NODES + 0.5 * (left + right)
        weights_in_hours = 0.5 * (right - left) * _GAUSS_WEIGHTS / 60.0
        for future_min, weight in zip(nodes, weights_in_hours, strict=True):
            waits = np.maximum(0.0, busy_until - future_min)
            waits[activation_min > time_min + EPS] = np.inf
            same_day = int((future_min + EPS) // MINUTES_PER_DAY) == int(
                (time_min + EPS) // MINUTES_PER_DAY
            )
            if same_day:
                capped = day_count >= DAILY_CAP
                waits[capped] = np.maximum(_next_midnight(time_min), busy_until[capped]) - future_min
            baseline = travel + waits[:, None]
            baseline_min = baseline.min(axis=0)
            multiplier = 1.0 if rate_multiplier is None else rate_multiplier(float(future_min))
            rates = data.demand * float(intraday_density(future_min / 60.0)) * multiplier
            excluded = np.broadcast_to(baseline, (len(candidate_list), *baseline.shape)).copy()
            excluded[np.arange(len(candidate_list)), candidate_indices, :] = np.inf
            other_min = excluded.min(axis=1)
            counter_busy = np.full(len(candidate_list), time_min + BUSY_MINUTES)
            counter_wait = np.maximum(0.0, counter_busy - future_min)
            if same_day:
                counter_capped = day_count[candidate_indices] + 1 >= DAILY_CAP
                counter_wait[counter_capped] = (
                    np.maximum(_next_midnight(time_min), counter_busy[counter_capped]) - future_min
                )
            counter_response = travel[candidate_indices, :] + counter_wait[:, None]
            after_min = np.minimum(other_min, counter_response)
            increases = np.maximum(0.0, after_min - baseline_min[None, :])
            increments = weight * (increases @ rates)
            for candidate, increment in zip(candidate_list, increments, strict=True):
                totals[candidate.ambulance_id] += float(increment)
    return {ambulance_id: max(0.0, value) for ambulance_id, value in totals.items()}


def _choose_a(data: ProblemData, fleet: list[Ambulance], call: Call, time_min: float) -> tuple[Ambulance | None, float | None]:
    candidates = _current_candidates(fleet, time_min)
    if not candidates:
        return None, None
    selected = min(
        candidates,
        key=lambda a: (data.distance[call.zone, a.site], a.day_count, a.site, a.ambulance_id),
    )
    return selected, None


def _choose_b(
    data: ProblemData,
    fleet: list[Ambulance],
    call: Call,
    time_min: float,
    beta: float,
    delta: float,
    rate_multiplier: Callable[[float], np.ndarray] | None = None,
) -> tuple[Ambulance | None, float | None]:
    candidates = _current_candidates(fleet, time_min)
    if not candidates:
        return None, None
    travel = {a.ambulance_id: 60.0 * data.distance[call.zone, a.site] / SPEED_KMH for a in candidates}
    nearest = min(travel.values())
    eligible = [a for a in candidates if travel[a.ambulance_id] - nearest <= delta + EPS]
    c_losses = cumulative_response_losses(
        data,
        fleet,
        eligible,
        time_min,
        rate_multiplier=rate_multiplier,
    )
    scored: list[tuple[float, float, int, int, Ambulance, float]] = []
    for ambulance in eligible:
        delta_t = travel[ambulance.ambulance_id] - nearest
        c_loss = c_losses[ambulance.ambulance_id]
        workload = (ambulance.day_count + 1) / DAILY_CAP
        score = delta_t + c_loss + beta * workload
        scored.append((score, delta_t, ambulance.day_count, ambulance.ambulance_id, ambulance, c_loss))
    selected = min(scored, key=lambda item: item[:4])
    return selected[4], selected[5]


def _earliest_regular_response(
    data: ProblemData,
    fleet: list[Ambulance],
    call: Call,
    time_min: float,
) -> float:
    values = []
    for ambulance in fleet:
        if ambulance.reserve:
            continue
        wait_from_now = _known_wait(ambulance, time_min, time_min)
        values.append(
            time_min
            - call.arrival_min
            + wait_from_now
            + PREP_MINUTES
            + 60.0 * data.distance[call.zone, ambulance.site] / SPEED_KMH
        )
    return min(values)


def _choose_c(
    data: ProblemData,
    fleet: list[Ambulance],
    call: Call,
    time_min: float,
    tau: float,
) -> tuple[Ambulance | None, float | None]:
    candidates = _current_candidates(fleet, time_min)
    if not candidates:
        return None, None
    regular = [a for a in candidates if not a.reserve]
    reserve = [a for a in candidates if a.reserve]
    predicted_regular = _earliest_regular_response(data, fleet, call, time_min)

    if predicted_regular > tau + EPS and reserve:
        fastest_reserve = min(
            reserve,
            key=lambda a: (data.distance[call.zone, a.site], a.day_count, a.site, a.ambulance_id),
        )
        immediate = (
            time_min
            - call.arrival_min
            + PREP_MINUTES
            + 60.0 * data.distance[call.zone, fastest_reserve.site] / SPEED_KMH
        )
        if immediate + EPS < predicted_regular:
            return fastest_reserve, None
    if not regular:
        return None, None
    selected = min(
        regular,
        key=lambda a: (data.distance[call.zone, a.site], a.day_count, a.site, a.ambulance_id),
    )
    return selected, None


def simulate(
    data: ProblemData,
    calls: list[Call],
    strategy: str,
    beta: float = 1.0,
    delta: float = 1.0,
    reserve_vector: Iterable[int] | None = None,
    tau: float = 5.0,
    rate_multiplier: Callable[[float], np.ndarray] | None = None,
    rate_multiplier_active_from: float | None = None,
    external_sites: Iterable[int] | None = None,
    external_activation_min: float = 0.0,
) -> tuple[pd.DataFrame, dict[str, float | int]]:
    fleet = build_fleet(
        data,
        reserve_vector if strategy == "C" else None,
        external_sites=external_sites,
        external_activation_min=external_activation_min,
    )
    queue: deque[Call] = deque()
    records: list[DispatchRecord] = []
    intervals: dict[int, list[tuple[float, float]]] = {a.ambulance_id: [] for a in fleet}
    daily_dispatches: Counter[tuple[int, int]] = Counter()
    dispatch_sequence: list[int] = []
    call_index = 0
    time_min = 0.0
    next_midnight = MINUTES_PER_DAY
    max_queue = 0

    while True:
        all_idle = all(a.busy_until <= time_min + EPS for a in fleet)
        if call_index >= len(calls) and not queue and all_idle:
            break
        next_arrival = calls[call_index].arrival_min if call_index < len(calls) else math.inf
        next_completion = min(
            (a.busy_until for a in fleet if a.busy_until > time_min + EPS),
            default=math.inf,
        )
        next_activation = min(
            (a.activation_min for a in fleet if a.activation_min > time_min + EPS),
            default=math.inf,
        )
        event_time = min(next_arrival, next_completion, next_activation, next_midnight)
        if not math.isfinite(event_time):
            raise RuntimeError("Simulation event calendar became empty before completion")
        time_min = event_time

        if abs(time_min - next_midnight) <= EPS:
            for ambulance in fleet:
                ambulance.day_count = 0
            next_midnight += MINUTES_PER_DAY
        while call_index < len(calls) and calls[call_index].arrival_min <= time_min + EPS:
            queue.append(calls[call_index])
            call_index += 1
        max_queue = max(max_queue, len(queue))

        while queue:
            call = queue[0]
            if strategy == "A":
                ambulance, c_loss = _choose_a(data, fleet, call, time_min)
            elif strategy == "B":
                active_multiplier = rate_multiplier
                if rate_multiplier_active_from is not None and time_min < rate_multiplier_active_from - EPS:
                    active_multiplier = None
                ambulance, c_loss = _choose_b(
                    data,
                    fleet,
                    call,
                    time_min,
                    beta,
                    delta,
                    rate_multiplier=active_multiplier,
                )
            elif strategy == "C":
                ambulance, c_loss = _choose_c(data, fleet, call, time_min, tau)
            else:
                raise ValueError(f"Unknown strategy: {strategy}")
            if ambulance is None:
                break

            queue.popleft()
            wait = time_min - call.arrival_min
            response = wait + PREP_MINUTES + 60.0 * data.distance[call.zone, ambulance.site] / SPEED_KMH
            dispatch_day = int(math.floor((time_min + EPS) / MINUTES_PER_DAY))
            if ambulance.day_count >= DAILY_CAP:
                raise AssertionError("Daily dispatch cap violated before assignment")
            ambulance.day_count += 1
            daily_dispatches[(ambulance.ambulance_id, dispatch_day)] += 1
            intervals[ambulance.ambulance_id].append((time_min, time_min + BUSY_MINUTES))
            ambulance.busy_until = time_min + BUSY_MINUTES
            dispatch_sequence.append(call.call_id)
            records.append(
                DispatchRecord(
                    call_id=call.call_id,
                    zone=call.zone,
                    ambulance_id=ambulance.ambulance_id,
                    site=ambulance.site,
                    arrival_min=call.arrival_min,
                    dispatch_min=time_min,
                    response_min=response,
                    wait_min=wait,
                    dispatch_day=dispatch_day,
                    c_loss_min=c_loss,
                )
            )
            max_queue = max(max_queue, len(queue))

    if len(records) != len(calls) or len({record.call_id for record in records}) != len(calls):
        raise AssertionError("Every call must be dispatched exactly once")
    if dispatch_sequence != sorted(dispatch_sequence):
        raise AssertionError("FCFS dispatch order was violated")
    if daily_dispatches and max(daily_dispatches.values()) > DAILY_CAP:
        raise AssertionError("Daily dispatch cap was exceeded")
    for ambulance_intervals in intervals.values():
        for (_, end), (next_start, _) in zip(ambulance_intervals, ambulance_intervals[1:]):
            if next_start < end - EPS:
                raise AssertionError("Ambulance busy intervals overlap")

    frame = pd.DataFrame([record.__dict__ for record in records]).sort_values("call_id")
    responses = frame["response_min"].to_numpy(dtype=float)
    waits = frame["wait_min"].to_numpy(dtype=float)
    delay_costs = np.asarray(delay_penalty_cost(responses), dtype=float)
    metrics: dict[str, float | int] = {
        "calls": int(len(frame)),
        "mean_response_min": float(np.mean(responses)),
        "strict_4min_rate": float(np.mean(responses <= 4.0 + EPS)),
        "p90_response_min": float(np.quantile(responses, 0.90)),
        "p95_response_min": float(np.quantile(responses, 0.95)),
        "mean_wait_min": float(np.mean(waits)),
        "max_wait_min": float(np.max(waits)),
        "max_queue": int(max_queue),
        "mean_delay_penalty_yuan_per_call": float(np.mean(delay_costs)),
        "total_delay_penalty_yuan": float(np.sum(delay_costs)),
        "max_daily_dispatches_per_ambulance": int(max(daily_dispatches.values(), default=0)),
        "simulation_end_min": float(time_min),
    }
    if strategy == "B":
        c_values = frame["c_loss_min"].dropna().to_numpy(dtype=float)
        metrics["mean_c_loss_min"] = float(np.mean(c_values))
        metrics["max_c_loss_min"] = float(np.max(c_values))
        metrics["min_c_loss_min"] = float(np.min(c_values))
    return frame, metrics


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def run_p1(project_root: Path, seed: int, days: int) -> Path:
    input_path = problem_statement_path(project_root)
    data = read_problem(input_path)
    q1 = solve_q1(data)
    calls = generate_calls(data, days=days, seed=seed)
    if not calls:
        raise RuntimeError("P1 call stream is empty")
    daily_call_counts = np.bincount(
        [int(call.arrival_min // MINUTES_PER_DAY) for call in calls],
        minlength=days,
    )
    if not np.all(daily_call_counts == DAILY_CALLS):
        raise AssertionError("Conditional NHPP must generate exactly 140 calls per day")

    output_dir = project_root / "results" / "p1"
    output_dir.mkdir(parents=True, exist_ok=True)
    assignment = pd.DataFrame(q1["assignment"], index=data.zone_ids, columns=data.site_ids)
    assignment.index.name = "zone_id"
    assignment.to_csv(output_dir / "q1_assignment.csv", encoding="utf-8-sig")

    call_frame = pd.DataFrame([call.__dict__ for call in calls])
    call_frame["zone_id"] = call_frame["zone"].map(lambda value: data.zone_ids[value])
    call_frame.to_csv(output_dir / "common_calls.csv", index=False, encoding="utf-8-sig")

    strategy_settings = {
        "A": {},
        "B": {"beta": 1.0, "delta": 1.0},
        "C": {"reserve_vector": [1, 0, 0, 0, 0, 0], "tau": 5.0},
    }
    metric_rows = []
    for strategy, settings in strategy_settings.items():
        records, metrics = simulate(data, calls, strategy=strategy, **settings)
        records.to_csv(output_dir / f"q2_records_{strategy}.csv", index=False, encoding="utf-8-sig")
        metric_rows.append({"strategy": strategy, **metrics})
    metrics_frame = pd.DataFrame(metric_rows)
    metrics_frame.to_csv(output_dir / "q2_metrics.csv", index=False, encoding="utf-8-sig")

    summary = {
        "mode": "P1_MINIMAL_SLICE",
        "input": {"path": str(input_path), "sha256": sha256(input_path)},
        "seed": seed,
        "days": days,
        "arrival_contract": {
            "model": "conditional NHPP with periodic double-Gaussian intraday density",
            "fixed_daily_calls": DAILY_CALLS,
            "daily_call_counts": daily_call_counts.tolist(),
            "zone_marking_probability": (data.demand / data.demand.sum()).tolist(),
        },
        "delay_penalty_contract": {
            "golden_response_minutes": GOLDEN_RESPONSE_MINUTES,
            "yuan_per_excess_minute_per_call": DELAY_PENALTY_YUAN_PER_MINUTE,
            "formula": "200 * max(response_min - 4, 0)",
        },
        "q1": {
            key: value.tolist() if isinstance(value, np.ndarray) else value
            for key, value in q1.items()
            if key != "assignment"
        },
        "q2": {
            "generated_calls": len(calls),
            "settings": strategy_settings,
            "metrics": metric_rows,
            "warning": "P1 smoke results are not tuned or final task-two conclusions.",
        },
    }
    summary_path = output_dir / "summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="A-problem optimization and dispatch simulation")
    subparsers = parser.add_subparsers(dest="command", required=True)
    p1 = subparsers.add_parser("p1", help="run the minimal P1 vertical slice")
    p1.add_argument("--project-root", type=Path, default=solution_root())
    p1.add_argument("--seed", type=int, default=20260811)
    p1.add_argument("--days", type=int, default=3)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.command == "p1":
        summary_path = run_p1(args.project_root.resolve(), seed=args.seed, days=args.days)
        print(summary_path)


if __name__ == "__main__":
    main()
