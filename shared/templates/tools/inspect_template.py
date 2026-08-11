# -*- coding: utf-8 -*-
"""解析国赛模板与格式要求文档的排版细节。"""
from docx import Document
from docx.shared import Pt

TMPL = r"D:\大二上学习资料\数模国赛\模板资料"


def dump_text(path, max_paras=200):
    print(f"\n===== {path.split(chr(92))[-1]} 正文 =====")
    doc = Document(path)
    for i, p in enumerate(doc.paragraphs[:max_paras]):
        t = p.text.strip()
        if t:
            print(f"[{i}] ({p.style.name}) {t[:100]}")


def dump_styles(path):
    print(f"\n===== {path.split(chr(92))[-1]} 样式 =====")
    doc = Document(path)
    sec = doc.sections[0]
    print(f"页面: {sec.page_width.cm:.2f}x{sec.page_height.cm:.2f} cm, "
          f"边距 上{sec.top_margin.cm:.2f} 下{sec.bottom_margin.cm:.2f} "
          f"左{sec.left_margin.cm:.2f} 右{sec.right_margin.cm:.2f}")
    for s in doc.styles:
        try:
            f = s.font
            pf = s.paragraph_format
            info = (f"{s.name} | font={f.name} size={f.size.pt if f.size else None} "
                    f"bold={f.bold} color={f.color.rgb if f.color and f.color.rgb else None} | "
                    f"before={pf.space_before.pt if pf.space_before else None} "
                    f"after={pf.space_after.pt if pf.space_after else None} "
                    f"line={pf.line_spacing}")
            if f.name or f.size or pf.space_before or pf.space_after:
                print(info)
        except Exception:
            pass


def dump_first_paras(path, n=45):
    print(f"\n===== {path.split(chr(92))[-1]} 前{n}段(含run字体) =====")
    doc = Document(path)
    for i, p in enumerate(doc.paragraphs[:n]):
        runs = [(r.text[:22], r.font.name, r.font.size.pt if r.font.size else None,
                 r.bold, r.font.italic) for r in p.runs[:2]]
        print(f"[{i}] style={p.style.name} align={p.alignment} | {p.text[:60]!r} | runs={runs}")


if __name__ == "__main__":
    dump_text(f"{TMPL}\\format2024.docx")
    dump_styles(f"{TMPL}\\2026数学建模国赛标准论文Word模板.docx")
    dump_first_paras(f"{TMPL}\\2026数学建模国赛标准论文Word模板.docx")
