"""Regression tests for stable repository-wide DOCX layout invariants."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPOSITORY_ROOT / "shared"))

from paper_format import apply_profile, load_profile, validate_docx_layout  # noqa: E402


def _pandoc_like_document():
    document = Document()
    title = document.add_paragraph("测试题名", style="Heading 1")
    major = document.add_paragraph("一、问题分析", style="Heading 2")
    secondary = document.add_paragraph("1.1 建模方法", style="Heading 3")
    tertiary = document.add_paragraph("1.1.1 数值方法", style="Heading 4")
    body = document.add_paragraph("正文")
    return document, title, major, secondary, tertiary, body


def test_pandoc_hierarchy_is_normalized_before_style_mapping() -> None:
    document, title, major, secondary, tertiary, _ = _pandoc_like_document()

    apply_profile(document)

    assert title.style.name == "Title"
    assert major.style.name == "Heading 1"
    assert secondary.style.name == "Heading 2"
    assert tertiary.style.name == "Heading 3"
    assert major.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert secondary.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert tertiary.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_layout_validator_enforces_portrait_headings_and_body_spacing() -> None:
    document, _, _, secondary, tertiary, body = _pandoc_like_document()
    profile = load_profile()
    apply_profile(document, profile)

    assert validate_docx_layout(document, profile) == []

    secondary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tertiary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.paragraph_format.line_spacing = 1.0
    document.sections[0].orientation = WD_ORIENT.LANDSCAPE
    document.sections[0].page_width, document.sections[0].page_height = (
        document.sections[0].page_height,
        document.sections[0].page_width,
    )

    errors = validate_docx_layout(document, profile)
    assert any("Heading 2" in error for error in errors)
    assert any("Heading 3" in error for error in errors)
    assert any("正文行距" in error for error in errors)
    assert any("page_width" in error or "page_height" in error for error in errors)

