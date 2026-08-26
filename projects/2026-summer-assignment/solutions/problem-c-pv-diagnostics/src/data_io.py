"""Authoritative DOCX input parsing and small serialization helpers for Problem C."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Mapping

import numpy as np
from docx import Document


EXPECTED_IDS = tuple(f"PV{i:03d}" for i in range(1, 101))


@dataclass(frozen=True)
class ProblemData:
    component_ids: tuple[str, ...]
    parameters: tuple[dict[str, Any], ...]
    generation: np.ndarray
    deviation_pct: np.ndarray
    reference_labels: tuple[str, ...]
    weather: tuple[dict[str, Any], ...]

    @property
    def historical_weather(self) -> tuple[dict[str, Any], ...]:
        return self.weather[:15]

    @property
    def day16_weather(self) -> dict[str, Any]:
        return self.weather[15]

    @property
    def station_generation(self) -> np.ndarray:
        return self.generation.sum(axis=0)


def _cell_texts(row: Any) -> list[str]:
    return [cell.text.strip() for cell in row.cells]


def load_problem_data(supporting_docx: Path | str) -> ProblemData:
    """Read the three complete supporting-data tables and enforce the shared contract."""

    path = Path(supporting_docx)
    if not path.is_file():
        raise FileNotFoundError(path)
    document = Document(path)
    if len(document.tables) != 3:
        raise ValueError(f"expected 3 supporting tables, found {len(document.tables)}")

    parameters: list[dict[str, Any]] = []
    for row in document.tables[0].rows[1:]:
        cells = _cell_texts(row)
        parameters.append(
            {
                "component_id": cells[0],
                "string": cells[1],
                "azimuth_deg": float(cells[2]),
                "tilt_deg": float(cells[3]),
                "rated_power_w": float(cells[4]),
                "service_years": float(cells[5].removesuffix("年")),
            }
        )

    component_ids: list[str] = []
    generation_rows: list[list[float]] = []
    deviation_pct: list[float] = []
    reference_labels: list[str] = []
    for row in document.tables[1].rows[1:]:
        cells = _cell_texts(row)
        component_ids.append(cells[0])
        generation_rows.append([float(value) for value in cells[1:16]])
        deviation_pct.append(float(cells[16]))
        reference_labels.append(cells[17])

    weather: list[dict[str, Any]] = []
    for day_index, row in enumerate(document.tables[2].rows[1:], start=1):
        cells = _cell_texts(row)
        weather.append(
            {
                "day": day_index,
                "day_label": cells[0],
                "irradiation_kwh_m2": float(cells[1]),
                "temperature_c": float(cells[2]),
                "wind_m_s": float(cells[3]),
                "weather_type": cells[4],
                "is_forecast": day_index == 16,
            }
        )

    parameter_ids = tuple(row["component_id"] for row in parameters)
    ids = tuple(component_ids)
    if ids != EXPECTED_IDS or parameter_ids != EXPECTED_IDS:
        raise ValueError("component IDs must be unique, aligned, and exactly PV001-PV100")
    if len(weather) != 16:
        raise ValueError(f"expected 16 weather rows, found {len(weather)}")

    generation = np.asarray(generation_rows, dtype=float)
    deviations = np.asarray(deviation_pct, dtype=float)
    if generation.shape != (100, 15):
        raise ValueError(f"expected generation shape (100, 15), found {generation.shape}")
    if deviations.shape != (100,):
        raise ValueError(f"expected 100 deviations, found {deviations.shape}")
    if not np.isfinite(generation).all() or not np.isfinite(deviations).all():
        raise ValueError("generation and deviation inputs must be finite")
    if any(not np.isfinite(row[key]) for row in weather for key in (
        "irradiation_kwh_m2", "temperature_c", "wind_m_s"
    )):
        raise ValueError("weather inputs must be finite")

    return ProblemData(
        component_ids=ids,
        parameters=tuple(parameters),
        generation=generation,
        deviation_pct=deviations,
        reference_labels=tuple(reference_labels),
        weather=tuple(weather),
    )


def write_csv(path: Path | str, rows: Iterable[Mapping[str, Any]], fieldnames: list[str] | None = None) -> None:
    rows = list(rows)
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if fieldnames is None:
        if not rows:
            raise ValueError(f"fieldnames required for empty CSV: {path}")
        fieldnames = list(rows[0].keys())
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def _json_default(value: Any) -> Any:
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, np.generic):
        return value.item()
    if isinstance(value, Path):
        return str(value)
    raise TypeError(f"not JSON serializable: {type(value).__name__}")


def write_json(path: Path | str, payload: Any) -> None:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2, default=_json_default)
        handle.write("\n")
