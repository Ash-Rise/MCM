import importlib.util
import json
import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = PROJECT_ROOT / "paper" / "paper.docx"
MARKDOWN_PATH = PROJECT_ROOT / "paper" / "paper.md"
POSTPROCESS_PATH = PROJECT_ROOT / "src" / "postprocess_paper_docx.py"
TABLE_BASELINE_PATH = PROJECT_ROOT / "paper" / "paper.docx"


def _load_postprocessor():
    spec = importlib.util.spec_from_file_location("postprocess_paper_docx", POSTPROCESS_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load {POSTPROCESS_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _style_num_id(style):
    p_pr = style.element.pPr
    if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
        return None
    return p_pr.numPr.numId.val


def _page_field_count(footer):
    return len(footer._element.findall(".//" + qn("w:instrText")))


def _east_asia_font(run):
    return run._r.get_or_add_rPr().get_or_add_rFonts().get(qn("w:eastAsia"))


def _table_layout_xml(table) -> bytes:
    """Return canonical table layout XML while ignoring mathematical payloads."""
    element = deepcopy(table._tbl)
    for math in element.xpath(".//m:oMath | .//m:oMathPara"):
        math.getparent().remove(math)
    for text in element.xpath(".//w:t"):
        text.text = ""
    volatile_attributes = {
        "paraId",
        "textId",
        "rsidR",
        "rsidRDefault",
        "rsidRPr",
        "rsidDel",
        "rsidP",
    }
    for node in element.iter():
        for attribute in list(node.attrib):
            if etree.QName(attribute).localname in volatile_attributes:
                del node.attrib[attribute]
    return etree.tostring(element, method="c14n")


def test_paper_typography_matches_reference_document():
    module = _load_postprocessor()
    document = Document()
    heading = document.add_heading("1 问题重述", level=2)
    body = document.add_paragraph("正文段落")
    caption = document.add_paragraph("图1 测试图")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格文字"

    module._format_document_typography(document)

    assert heading.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert heading.runs[0].font.size == Pt(14)
    assert heading.runs[0].bold is True
    assert _east_asia_font(heading.runs[0]) == "黑体"
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.style.name == "Body Text"
    assert body.paragraph_format.first_line_indent == Pt(24)
    assert body.paragraph_format.line_spacing == 1.5
    assert body.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
    assert body._p.pPr.find(qn("w:snapToGrid")).get(qn("w:val")) == "0"
    assert body.runs[0].font.size == Pt(12)
    assert _east_asia_font(body.runs[0]) == "宋体"
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert caption.runs[0].font.size == Pt(10.5)
    table_paragraph = table.cell(0, 0).paragraphs[0]
    assert table_paragraph.runs[0].font.size == Pt(10.5)
    assert table_paragraph.paragraph_format.line_spacing == 1.15


def test_typography_normalizes_dangling_pandoc_body_styles():
    module = _load_postprocessor()
    document = Document()
    first = document.add_paragraph("标题后的首段")
    first._p.get_or_add_pPr().get_or_add_pStyle().val = "FirstParagraph"
    listed = document.add_paragraph("编号正文")
    listed_p_pr = listed._p.get_or_add_pPr()
    listed_p_pr.get_or_add_pStyle().val = "Compact"
    listed_p_pr.get_or_add_numPr().get_or_add_numId().val = 1

    module._format_document_typography(document)

    assert first.style.name == "Body Text"
    assert listed.style.name == "Body Text"
    assert listed._p.pPr.numPr.numId.val == 1
    for paragraph in (first, listed):
        assert paragraph.paragraph_format.line_spacing == 1.5
        assert paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE


def test_caption_detection_does_not_center_narrative_sentences():
    module = _load_postprocessor()
    document = Document()
    caption = document.add_paragraph("表3 优化后的服务分配")
    narrative = document.add_paragraph("表3给出非零服务分配。")

    module._format_document_typography(document)

    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert narrative.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_document_grid_matches_reference_body_line_pitch():
    module = _load_postprocessor()
    document = Document()

    module._set_document_line_grid(document)

    document_grid = document.sections[0]._sectPr.find(qn("w:docGrid"))
    assert document_grid is not None
    assert document_grid.get(qn("w:type")) == "lines"
    assert document_grid.get(qn("w:linePitch")) == "360"
    assert document_grid.get(qn("w:charSpace")) == "0"


def test_disable_grid_places_snap_before_spacing_and_alignment():
    module = _load_postprocessor()
    document = Document()
    paragraph = document.add_paragraph("正文")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    module._disable_paragraph_grid(paragraph)

    child_tags = [child.tag for child in paragraph._p.pPr]
    snap_index = child_tags.index(qn("w:snapToGrid"))
    assert snap_index < child_tags.index(qn("w:spacing"))
    assert snap_index < child_tags.index(qn("w:jc"))


def test_markdown_math_avoids_bare_star_superscripts():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")

    assert not re.search(r"\^\*", markdown)


def test_title_and_abstract_meet_frozen_review_contract():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    title, remainder = markdown.split("\n", maxsplit=1)
    abstract = remainder.split("## 摘 要", maxsplit=1)[1].split(
        "## 一、问题重述", maxsplit=1
    )[0]
    abstract_paragraphs = [
        paragraph.strip()
        for paragraph in re.split(r"\n\s*\n", abstract)
        if paragraph.strip() and not paragraph.startswith("**关键词：**")
    ]

    assert title == "# 基于容量约束运输规划与条件NHPP仿真的急救车辆配置调度优化"
    assert len(title.removeprefix("# ")) <= 32
    assert len(abstract_paragraphs) == 5
    assert all(f"针对任务{number}" in abstract for number in "一二三")
    assert "求解器状态为最优" in abstract
    assert "共同随机数" in abstract
    assert "适用边界" in abstract
    assert "$" not in abstract


def test_markdown_uses_chinese_top_level_headings_and_superscript_citations():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    expected_headings = (
        "## 一、问题重述",
        "## 二、问题分析",
        "## 三、模型假设",
        "## 四、符号说明",
        "## 五、任务一：容量约束站点与服务分配",
        "## 六、任务二：连续多日随机呼叫与车辆调度",
        "## 七、任务三：连续事故时长下的应急调度与效果评估",
        "## 八、模型评价",
    )

    for heading in expected_headings:
        assert heading in markdown
    assert not re.search(r"^## [1-8] ", markdown, flags=re.MULTILINE)
    assert len(re.findall(r"\^\\\[[1-4]\\\]\^", markdown)) == 4
    body, references = markdown.split("## 参考文献", maxsplit=1)
    assert not re.search(r"(?<!\\)\[[1-4]\]", body)
    assert all(f"[{index}]" in references for index in range(1, 5))


def test_markdown_uses_standard_periodic_and_exponential_notation():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    body = markdown.split("## 参考文献", maxsplit=1)[0]

    assert r"\bmod" not in body
    assert r"\exp" not in body
    assert body.count(r"f(t\%24)") == 1
    assert r"\mathrm e^{-\frac{(t-\mu+24k)^2}{2\sigma^2}}" in body


def test_arrival_rate_term_normalization_preserves_run_structure():
    module = _load_postprocessor()
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("纵轴为严格4")
    paragraph.add_run(" ")
    last = paragraph.add_run("min响应率，继续说明。")
    first.bold = True
    last.italic = True

    assert module.normalize_arrival_rate_terms(document) == 1
    assert paragraph.text == "纵轴为4分钟内到达率，继续说明。"
    assert len(paragraph.runs) == 3
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[2].italic is True


def test_markdown_has_algorithm_design_for_all_three_tasks():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")

    assert markdown.count("算法设计") == 3
    assert "### 5.5 算法设计" in markdown
    assert "### 6.7 算法设计" in markdown
    assert "### 7.6 算法设计" in markdown
    assert markdown.count("| Step ") == 21
    assert "表2 任务一混合整数线性规划模型" in markdown
    assert "表4 任务一算法步骤" in markdown
    assert "表6 任务二算法步骤" in markdown
    assert "表10 任务三算法步骤" in markdown


def test_figure_numbers_are_continuous_and_appendix_matches():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    captions = [
        int(number)
        for number in re.findall(r"^图(\d+)\s", markdown, flags=re.MULTILINE)
    ]

    assert captions == list(range(1, 12))
    assert "图1至图11均由" in markdown


def test_task_three_separates_emergency_dispatch_from_external_support():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    task_three = markdown.split("## 七、任务三", 1)[1].split("## 八、模型评价", 1)[0]
    section_73 = task_three.split("### 7.3 应急响应方案", 1)[1].split("### 7.4", 1)[0]

    assert "10个事故区域、10个时长节点、10个随机种子和7种外援数量" in task_three
    assert "### 7.3 应急响应方案" in task_three
    assert "### 7.9" not in task_three
    assert "余量仅为4次/日，即2.86%" in task_three
    assert "只改变候选车辆的派出顺序" in task_three
    assert "这里的$B_N$仍包含事故呼叫" in task_three
    assert "$B_N$可作为$B_E$在同一事故情景下的对照" in task_three
    assert "共同随机数控制了事故规模" in task_three
    assert r"D_{k,s}(H)=T_{k,s}^{B_E}(H)-T_{k,s}^{B_N}(H)" in task_three
    assert "事故新增需求和派车调整" in task_three
    assert "识别事故信息修正的作用" in task_three
    assert "评价扩充运力带来的改善" in task_three
    assert "持续排队" not in task_three
    assert "在$H=6,8,10,11,12$ h五个长时节点" in task_three
    assert "六个站点和12辆车保持任务一的配置" in section_73
    assert "外援" not in section_73
    assert "尚未完成的事故期呼叫继续处理" in section_73
    assert "3辆是90%响应改善达成率规则下的最小统一外援数" in task_three
    assert "5辆是满足该目标的最小统一数量" in task_three
    assert "1辆方案的单位车辆收益最高" in task_three
    assert "3辆方案在长时事故下综合性价比最高" in task_three
    assert "经济最优数量仍由具体调配成本决定" not in task_three
    assert r"c_m(H)" not in task_three
    assert "表13" not in task_three
    assert "第1～6辆外援的$\\Delta P_m$依次为39.01、19.69、8.24、2.48、0.76和0.19万元" in task_three
    assert r"\Delta P_m(H)=P_{m-1}(H)-P_m(H)" in task_three
    for obsolete_symbol in ("G_m(H)", "g_m(H)", "S_m(H)", "A_m(H)", "B_m(H)"):
        assert obsolete_symbol not in task_three
    assert r"\bar M_k(H)" not in task_three
    assert r"s_M(H)" not in task_three
    assert "获取得分" not in task_three


def test_figure_eleven_includes_response_gain_and_marginal_penalty_panels():
    source = (PROJECT_ROOT / "src" / "generate_figures.py").read_text(encoding="utf-8")

    assert 'long_durations = [6.0, 8.0, 10.0, 11.0, 12.0]' in source
    assert 'ax_share.set_ylabel("响应改善达成率（%）")' in source
    assert 'ax_share.set_title("长时事故的响应改善达成率"' in source
    assert '"6辆外援=100%"' in source
    assert 'values="marginal_break_even_cost_yuan_mean"' in source
    assert "colors.LogNorm" in source
    assert 'ax_penalty.set_title("新增第m辆外援的边际避免罚金"' in source
    assert '"边际避免罚金（万元/辆·事故情景，对数色阶）"' in source


def test_task_one_uses_only_planning_coverage_and_compact_capacity_proof():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    task_one = markdown.split("## 五、任务一", 1)[1].split("## 六、任务二", 1)[0]

    assert "表3 各建站规模运力可行性" in task_one
    assert "| 最大配车/辆 | 3 | 5 | 7 | 9 | 11 | 12 |" in task_one
    assert "严格4 min中心代理" not in task_one
    assert "0.75 km" not in task_one
    assert "60.714%" not in task_one
    assert "3 km规划服务覆盖率为86.429%" in task_one


def test_strategy_c_defines_dispatch_selection_and_out_of_sample_check():
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    strategy_c = markdown.split("### 6.5 策略C", 1)[1].split("### 6.6", 1)[0]

    assert "47种非零备用向量" in strategy_c
    assert "235个候选方案" in strategy_c
    assert r"\widehat T_e^{\mathrm{reg}}>\tau" in strategy_c
    assert r"T_e^R<\widehat T_e^{\mathrm{reg}}" in strategy_c
    assert "严格词典序" in strategy_c
    assert r"\boldsymbol r^{\ast}=(0,0,1,0,0,0)" in strategy_c
    assert r"\tau^{\ast}=7" in strategy_c

    task_two_results = markdown.split("### 6.8 最终结果与方案选择", 1)[1].split(
        "### 6.9", 1
    )[0]
    assert r"U_{0.95}" in task_two_results
    assert "$U_{0.95}=-0.0091$ min" in task_two_results


def test_omml_normalizer_fixes_pandoc_matrix_child_order():
    module = _load_postprocessor()
    document = Document()
    paragraph = document.add_paragraph()
    run_properties = OxmlElement("m:rPr")
    run_properties.append(OxmlElement("m:nor"))
    run_properties.append(OxmlElement("m:scr"))
    run_properties.append(OxmlElement("m:sty"))
    paragraph._p.append(run_properties)
    column_properties = OxmlElement("m:mcPr")
    column_properties.append(OxmlElement("m:mcJc"))
    column_properties.append(OxmlElement("m:count"))
    paragraph._p.append(column_properties)

    assert module._normalize_omml_matrix_properties(document) == (1, 1)
    assert [node.tag for node in run_properties] == [qn("m:nor")]
    assert [node.tag for node in column_properties] == [qn("m:count"), qn("m:mcJc")]


def test_postprocessor_removes_duplicate_bookmark_ends_imported_with_locked_table(tmp_path):
    module = _load_postprocessor()
    source = Document()
    source.add_paragraph("测试题名")
    source.add_paragraph("一、问题重述")
    body = source.add_paragraph("正文")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "25")
    bookmark_start.set(qn("w:name"), "body")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "25")
    body._p.append(bookmark_start)
    body._p.append(bookmark_end)
    source.add_table(rows=1, cols=1).cell(0, 0).text = "相同内容"

    baseline = Document()
    baseline.add_table(rows=1, cols=1).cell(0, 0).text = "相同内容"
    duplicate_end = OxmlElement("w:bookmarkEnd")
    duplicate_end.set(qn("w:id"), "25")
    baseline.tables[0]._tbl.append(duplicate_end)

    source_path = tmp_path / "source.docx"
    baseline_path = tmp_path / "baseline.docx"
    output_path = tmp_path / "output.docx"
    source.save(source_path)
    baseline.save(baseline_path)
    module.postprocess_docx(source_path, output_path, table_baseline=baseline_path)

    output = Document(output_path)
    bookmark_ends = output.element.body.findall(".//" + qn("w:bookmarkEnd"))
    assert [node.get(qn("w:id")) for node in bookmark_ends].count("25") == 1


def test_postprocessor_preserves_omml_imported_with_locked_table(tmp_path):
    module = _load_postprocessor()
    source = Document()
    source.add_paragraph("测试题名")
    source.add_paragraph("一、问题重述")
    source.add_table(rows=1, cols=1).cell(0, 0).text = "相同内容"

    baseline = Document()
    baseline.add_table(rows=1, cols=1).cell(0, 0).text = "相同内容"
    run_properties = OxmlElement("m:rPr")
    run_properties.append(OxmlElement("m:sty"))
    run_properties.append(OxmlElement("m:scr"))
    baseline.tables[0].cell(0, 0).paragraphs[0]._p.append(run_properties)

    source_path = tmp_path / "source.docx"
    baseline_path = tmp_path / "baseline.docx"
    output_path = tmp_path / "output.docx"
    source.save(source_path)
    baseline.save(baseline_path)
    module.postprocess_docx(source_path, output_path, table_baseline=baseline_path)

    output = Document(output_path)
    output_properties = output.element.body.find(".//" + qn("m:rPr"))
    assert output_properties is not None
    assert [node.tag for node in output_properties] == [qn("m:sty"), qn("m:scr")]


def test_complete_docx_postprocessor_removes_heading_numbering_and_fixes_tables(tmp_path):
    module = _load_postprocessor()
    output = tmp_path / "complete-paper.docx"
    source_document = Document(DOCX_PATH)
    module.postprocess_docx(
        DOCX_PATH,
        output,
        table_baseline=TABLE_BASELINE_PATH,
    )

    document = Document(output)
    table_baseline = Document(TABLE_BASELINE_PATH)
    assert document.paragraphs[0].style.name == "Title"
    assert document.paragraphs[0].runs[0].font.size == Pt(15)
    assert _style_num_id(document.styles["Heading 2"]) == 0
    assert _style_num_id(document.styles["Heading 3"]) == 0
    assert _page_field_count(document.sections[0].footer) == 1
    assert _page_field_count(document.sections[0].first_page_footer) == 1
    assert document.sections[0].footer.paragraphs[0].alignment == 1
    assert document.sections[0].first_page_footer.paragraphs[0].alignment == 1
    document_grid = document.sections[0]._sectPr.find(qn("w:docGrid"))
    assert document_grid is not None
    assert document_grid.get(qn("w:linePitch")) == "360"

    display_math_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if not paragraph.text.strip()
        and paragraph._p.findall(".//" + qn("m:oMath"))
    ]
    assert display_math_paragraphs
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in display_math_paragraphs
    )
    table_math_paragraphs = [
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph._p.findall(".//" + qn("m:oMath"))
    ]
    assert table_math_paragraphs
    allowed_left_math_ids = {
        id(paragraph._p)
        for table_index in (3, 5, 9)
        for row in document.tables[table_index].rows[1:]
        for column_index in (0, 1)
        for paragraph in row.cells[column_index].paragraphs
        if paragraph._p.findall(".//" + qn("m:oMath"))
    }
    allowed_left_math_ids.update(
        id(paragraph._p)
        for row in document.tables[0].rows[1:]
        for paragraph in row.cells[1].paragraphs
        if paragraph._p.findall(".//" + qn("m:oMath"))
    )
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        or id(paragraph._p) in allowed_left_math_ids
        for paragraph in table_math_paragraphs
    )
    horizontal_scripts = [
        script
        for script_tag in ("m:sub", "m:sup")
        for script in document.element.body.findall(".//" + qn(script_tag))
        if "".join(script.itertext())
        in {"ij", "ea", "i=1", "j=1", "i(e),j(a)", "wait", "resp"}
    ]
    assert horizontal_scripts
    assert all(len(script.findall(qn("m:r"))) == 1 for script in horizontal_scripts)

    for paragraph in document.paragraphs:
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        is_expected_center = (
            style_name in {"Title", "Heading 1", "Heading 2"}
            or module._is_figure_or_table_caption(text)
            or bool(paragraph._p.findall(".//" + qn("w:drawing")))
            or (not text and bool(paragraph._p.findall(".//" + qn("m:oMath"))))
        )
        assert is_expected_center, f"Unexpected centered paragraph: {text[:80]}"

    assert len(document.tables) == len(table_baseline.tables) == 12
    assert document.tables[0].cell(8, 2).text == "1/h，次/h"
    for table_index, (table, source_table, baseline_table) in enumerate(zip(
        document.tables,
        source_document.tables,
        table_baseline.tables,
        strict=True,
    )):
        assert module._table_text(table) == module._table_text(source_table)
        assert table._tbl.xml == baseline_table._tbl.xml
        if table_index == 6:
                widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
                assert widths[0] > max(widths[1:])
                assert max(widths[1:]) - min(widths[1:]) <= 1


def test_table_baseline_rejects_content_drift():
    module = _load_postprocessor()
    target = Document()
    target.add_table(rows=1, cols=1).cell(0, 0).text = "new content"
    baseline = Document()
    baseline.add_table(rows=1, cols=1).cell(0, 0).text = "old content"

    try:
        module._replace_tables_from_baseline(target, baseline)
    except ValueError as error:
        assert "Table 1 content differs" in str(error)
    else:
        raise AssertionError("A changed table was silently replaced by the baseline")


def test_table_layout_only_mode_retains_content_and_reports_drift():
    module = _load_postprocessor()
    target = Document()
    target_table = target.add_table(rows=1, cols=1)
    target_table.cell(0, 0).text = "new content"
    baseline = Document()
    baseline_table = baseline.add_table(rows=1, cols=1)
    baseline_table.cell(0, 0).text = "old content"
    baseline_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    assert module._copy_table_layout_from_baseline(target, baseline) == [1]
    assert target.tables[0].cell(0, 0).text == "new content"
    assert (
        target.tables[0].cell(0, 0).paragraphs[0].alignment
        == WD_ALIGN_PARAGRAPH.CENTER
    )


def test_hybrid_table_lock_preserves_unchanged_tables_and_reports_changed_tables():
    module = _load_postprocessor()
    target = Document()
    target.add_table(rows=1, cols=1).cell(0, 0).text = "unchanged"
    target.add_table(rows=1, cols=1).cell(0, 0).text = "new content"
    baseline = Document()
    baseline.add_table(rows=1, cols=1).cell(0, 0).text = "unchanged"
    baseline.tables[0].cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    baseline.add_table(rows=1, cols=2).cell(0, 0).text = "old content"

    assert module._replace_unchanged_tables_from_baseline(target, baseline) == [2]
    assert target.tables[0]._tbl.xml == baseline.tables[0]._tbl.xml
    assert target.tables[1].cell(0, 0).text == "new content"


def test_exact_table_lock_survives_docx_save_and_reload(tmp_path):
    module = _load_postprocessor()
    target = Document()
    target_table = target.add_table(rows=2, cols=2)
    baseline = Document()
    baseline_table = baseline.add_table(rows=2, cols=2)
    for row_index in range(2):
        for column_index in range(2):
            text = f"r{row_index}c{column_index}"
            target_table.cell(row_index, column_index).text = text
            baseline_table.cell(row_index, column_index).text = text
    baseline_table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    module._replace_tables_from_baseline(target, baseline)

    target_path = tmp_path / "target.docx"
    baseline_path = tmp_path / "baseline.docx"
    target.save(target_path)
    baseline.save(baseline_path)
    target = Document(target_path)
    baseline = Document(baseline_path)

    assert target.tables[0]._tbl.xml == baseline.tables[0]._tbl.xml
    assert target.tables[0].cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_manual_table_baseline_keeps_table8_execution_notes_left_aligned():
    baseline = Document(TABLE_BASELINE_PATH)
    execution_cells = [row.cells[3] for row in baseline.tables[7].rows[1:]]

    assert execution_cells
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT
        for cell in execution_cells
        for paragraph in cell.paragraphs
    )


def test_default_table_baseline_hash_is_frozen():
    module = _load_postprocessor()

    assert TABLE_BASELINE_PATH == module.DEFAULT_TABLE_BASELINE
    assert module.sha256_file(TABLE_BASELINE_PATH) == module.DEFAULT_TABLE_BASELINE_SHA256

def test_postprocessor_selects_complete_paper_table_geometry():
    module = _load_postprocessor()

    weights = module.table_width_weights_for_count(12)

    assert len(weights) == 12
    assert tuple(len(item) for item in weights) == (3, 3, 7, 2, 6, 2, 4, 4, 4, 2, 5, 5)


def test_postprocessor_rejects_unknown_table_count():
    module = _load_postprocessor()

    try:
        module.table_width_weights_for_count(10)
    except ValueError as error:
        assert "12 tables" in str(error)
    else:
        raise AssertionError("Unknown paper table count was accepted")


def test_rebind_conversion_manifest_tracks_postprocessed_output(tmp_path):
    module = _load_postprocessor()
    source = tmp_path / "paper.md"
    source.write_text("updated source", encoding="utf-8")
    output = tmp_path / "complete.docx"
    output.write_bytes(b"postprocessed-docx")
    manifest_path = tmp_path / "intermediate.conversion.json"
    manifest_path.write_text(
        json.dumps(
            {
                "source": str(source),
                "source_sha256": "stale-source",
                "project_sha256": "stale-project",
                "output": str(tmp_path / "intermediate.docx"),
                "output_sha256": "stale",
            }
        ),
        encoding="utf-8",
    )

    table_baseline = tmp_path / "table-baseline.docx"
    table_baseline.write_bytes(b"manual-table-layout")
    final_manifest = module.rebind_conversion_manifest(
        manifest_path, output, table_baseline=table_baseline
    )

    assert final_manifest == output.with_suffix(".conversion.json")
    data = json.loads(final_manifest.read_text(encoding="utf-8"))
    assert data["output"] == str(output.resolve())
    assert data["output_sha256"] == module.sha256_file(output)
    assert data["source_sha256"] == module.sha256_file(source)
    assert data["project_sha256"] == module.sha256_file(source)
    assert data["postprocess"]["tool"] == "src/postprocess_paper_docx.py"
    assert data["postprocess"]["table_baseline"] == str(table_baseline.resolve())
    assert data["postprocess"]["table_baseline_sha256"] == module.sha256_file(
        table_baseline
    )
    assert data["postprocess"]["table_baseline_mode"] == "complete_table_xml"


def test_body_image_alt_text_uses_following_figure_caption(tmp_path):
    module = _load_postprocessor()
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (4, 4), "white").save(image_path)
    document = Document()
    document.add_picture(str(image_path))
    document.add_paragraph("图1 测试图像")
    document.add_picture(str(image_path))
    document.add_paragraph("普通段落")

    assert module._set_body_image_alt_text(document) == 1
    properties = document.paragraphs[0]._p.find(".//" + qn("wp:docPr"))
    assert properties is not None
    assert properties.get("descr") == "图1 测试图像"
    assert properties.get("title") == "图1"
    untouched = document.paragraphs[2]._p.find(".//" + qn("wp:docPr"))
    assert untouched is not None
    assert untouched.get("descr") is None


def test_postprocessor_keeps_abstract_on_first_page():
    module = _load_postprocessor()
    document = Document()
    document.add_paragraph("摘要内容")
    body_heading = document.add_paragraph("一、问题重述")

    module._keep_abstract_on_first_page(document)

    assert body_heading.paragraph_format.page_break_before is True
