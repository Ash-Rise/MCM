"""Apply the repository paper profile to the generated Problem B DOCX."""

from __future__ import annotations

import argparse
import os
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TABLE_WIDTH_WEIGHTS = (
    (0.07, 0.13, 0.16, 0.11, 0.19, 0.34),
    (0.13, 0.67, 0.20),
    (0.08, 0.28, 0.10, 0.12, 0.14, 0.09, 0.09, 0.10),
    (0.09, 0.10, 0.25, 0.18, 0.20, 0.18),
    (0.40, 0.20, 0.20, 0.20),
)


def set_run_font(run, *, east_asia: str, size: float, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)


def remove_numbering(element) -> None:
    p_pr = element.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def remove_style_numbering(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def set_cell_margins(cell, *, top: int = 55, bottom: int = 55, side: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        following_tags = {
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
            qn("w:headers"),
            qn("w:cellIns"),
            qn("w:cellDel"),
            qn("w:cellMerge"),
            qn("w:tcPrChange"),
        }
        insertion_index = next(
            (index for index, child in enumerate(tc_pr) if child.tag in following_tags),
            len(tc_pr),
        )
        tc_pr.insert(insertion_index, tc_mar)
    for name, value in (("top", top), ("start", side), ("bottom", bottom), ("end", side)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, weights: tuple[float, ...], width_dxa: int) -> None:
    if len(table.columns) != len(weights):
        raise ValueError(f"Expected {len(weights)} columns, found {len(table.columns)}")
    widths = [round(width_dxa * weight / sum(weights)) for weight in weights]
    widths[-1] += width_dxa - sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_edge(cell, name: str, *, value: str, size: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        following_tags = {
            qn("w:shd"),
            qn("w:noWrap"),
            qn("w:tcMar"),
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
            qn("w:headers"),
            qn("w:cellIns"),
            qn("w:cellDel"),
            qn("w:cellMerge"),
            qn("w:tcPrChange"),
        }
        insertion_index = next(
            (index for index, child in enumerate(tc_pr) if child.tag in following_tags),
            len(tc_pr),
        )
        tc_pr.insert(insertion_index, borders)
    edge = borders.find(qn(f"w:{name}"))
    if edge is None:
        edge = OxmlElement(f"w:{name}")
        borders.append(edge)
    edge.set(qn("w:val"), value)
    if value == "single":
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), "000000")


def format_tables(document) -> None:
    if len(document.tables) != len(TABLE_WIDTH_WEIGHTS):
        raise ValueError(
            f"Expected {len(TABLE_WIDTH_WEIGHTS)} tables, found {len(document.tables)}"
        )
    section = document.sections[0]
    content_width = int((section.page_width - section.left_margin - section.right_margin) / 635)
    for table_index, (table, weights) in enumerate(
        zip(document.tables, TABLE_WIDTH_WEIGHTS, strict=True)
    ):
        set_table_geometry(table, weights, content_width - 120)
        for row_index, row in enumerate(table.rows):
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            if row_index == 0 and tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for edge_name in ("top", "start", "bottom", "end", "insideH", "insideV"):
                    set_edge(cell, edge_name, value="nil")
                if row_index == 0:
                    set_edge(cell, "top", value="single", size=10)
                    set_edge(cell, "bottom", value="single", size=4)
                if row_index == len(table.rows) - 1:
                    set_edge(cell, "bottom", value="single", size=10)
                for paragraph in cell.paragraphs:
                    p_style = paragraph._p.get_or_add_pPr().find(qn("w:pStyle"))
                    if p_style is not None:
                        paragraph._p.get_or_add_pPr().remove(p_style)
                    remove_numbering(paragraph._p)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.keep_with_next = (
                        table_index != 1 and row_index < len(table.rows) - 1
                    )
                    for run in paragraph.runs:
                        set_run_font(run, east_asia="宋体", size=10, bold=(row_index == 0))


def set_page_field(footer) -> None:
    element = footer._element
    for child in list(element):
        element.remove(child)
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    paragraph.append(p_pr)
    for kind, text in (("begin", None), (None, " PAGE "), ("separate", None)):
        run = OxmlElement("w:r")
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run.append(fld)
        else:
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = text
            run.append(instruction)
        paragraph.append(run)
    value_run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = "1"
    value_run.append(value)
    paragraph.append(value_run)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph.append(end_run)
    element.append(paragraph)


def format_document(document) -> None:
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        remove_style_numbering(style)

    title = document.paragraphs[0]
    remove_numbering(title._p)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.line_spacing = 1.0
    title.paragraph_format.space_after = Pt(3)
    for run in title.runs:
        set_run_font(run, east_asia="宋体", size=14.5, bold=True)

    for paragraph in document.paragraphs[1:]:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            remove_numbering(paragraph._p)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(2.5)
            paragraph.paragraph_format.space_after = Pt(2.5)
            if style_name in {"Heading 1", "Heading 2"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                font, size = "黑体", 14
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                font, size = "宋体", 12
            for run in paragraph.runs:
                set_run_font(run, east_asia=font, size=size, bold=True)
            continue
        if paragraph._p.findall(".//" + qn("w:drawing")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            continue
        if re.match(r"^[图表]\s*\d+\s", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.keep_with_next = text.startswith("表 ")
            for run in paragraph.runs:
                set_run_font(run, east_asia="宋体", size=10.5)
            continue
        if re.match(r"^\[\d+\]", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(21)
            paragraph.paragraph_format.first_line_indent = Pt(-21)
            paragraph.paragraph_format.line_spacing = 1.25
            for run in paragraph.runs:
                set_run_font(run, east_asia="宋体", size=10.5)
            continue
        if not text and paragraph._p.findall(".//" + qn("m:oMath")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        has_num = paragraph._p.get_or_add_pPr().find(qn("w:numPr")) is not None
        paragraph.paragraph_format.first_line_indent = None if has_num else Pt(24)
        if text.startswith("关键词"):
            paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            set_run_font(run, east_asia="宋体", size=12)

    first_body_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().endswith("问题重述")
    )
    first_body_heading.paragraph_format.page_break_before = True
    format_tables(document)
    for section in document.sections:
        set_page_field(section.footer)
        set_page_field(section.first_page_footer)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.input)
    format_document(document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        dir=args.output.parent, suffix=".docx", delete=False
    ) as temporary:
        temporary_path = Path(temporary.name)
    try:
        document.save(temporary_path)
        os.replace(temporary_path, args.output)
    finally:
        temporary_path.unlink(missing_ok=True)
if __name__ == "__main__":
    main()
