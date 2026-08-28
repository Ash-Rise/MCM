"""Regression checks for the final Problem C DOCX resource contract."""

from __future__ import annotations

import sys
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = PROJECT_ROOT.parents[3]
sys.path.insert(0, str(REPOSITORY_ROOT / "shared"))

from paper_format import load_profile, validate_docx_resources  # noqa: E402


DOCX_PATH = PROJECT_ROOT / "paper" / "paper.docx"


class ProblemCPaperDocxFormatTests(unittest.TestCase):
    def test_docx_has_profile_owned_theme_styles_and_font_table(self) -> None:
        document = Document(DOCX_PATH)
        self.assertEqual(validate_docx_resources(document, load_profile()), [])
        with ZipFile(DOCX_PATH) as archive:
            self.assertNotIn(b"Aptos", archive.read("word/theme/theme1.xml"))
            self.assertNotIn(b"Aptos", archive.read("word/fontTable.xml"))

    def test_docx_uses_the_approved_shared_page_and_paragraph_profile(self) -> None:
        profile = load_profile()
        document = Document(DOCX_PATH)
        section = document.sections[0]
        page = profile["page"]
        self.assertEqual(section.page_width.twips, page["width_twips"])
        self.assertEqual(section.page_height.twips, page["height_twips"])
        self.assertEqual(section.top_margin.twips, page["margins_twips"]["top"])
        self.assertEqual(section.header_distance.twips, page["header_distance_twips"])
        self.assertEqual(section.footer_distance.twips, page["footer_distance_twips"])
        self.assertTrue(section.different_first_page_header_footer)
        grid = section._sectPr.find(qn("w:docGrid"))
        self.assertEqual(grid.get(qn("w:type")), page["document_grid"]["type"])
        self.assertEqual(
            grid.get(qn("w:linePitch")),
            str(page["document_grid"]["line_pitch_twips"]),
        )
        self.assertEqual(
            grid.get(qn("w:charSpace")),
            str(page["document_grid"]["char_space_twips"]),
        )

        body = next(
            paragraph
            for paragraph in document.paragraphs[2:]
            if paragraph.text.strip() and not paragraph.style.name.startswith("Heading")
        )
        body_profile = profile["typography"]["body"]
        self.assertEqual(body.paragraph_format.line_spacing, body_profile["line_spacing_multiple"])
        self.assertEqual(body.paragraph_format.space_before, Pt(body_profile["space_before_pt"]))
        self.assertEqual(body.paragraph_format.space_after, Pt(body_profile["space_after_pt"]))
        self.assertEqual(body.paragraph_format.first_line_indent, Pt(body_profile["first_line_indent_pt"]))
        self.assertEqual(
            body._p.get_or_add_pPr().find(qn("w:snapToGrid")).get(qn("w:val")),
            "0",
        )

        first_body_heading = next(
            paragraph
            for paragraph in document.paragraphs[2:]
            if paragraph.style.name in {"Heading 1", "Heading 2"}
        )
        self.assertTrue(first_body_heading.paragraph_format.page_break_before)

        table_paragraph = document.tables[0].cell(0, 0).paragraphs[0]
        table_profile = profile["typography"]["table_text"]
        self.assertEqual(
            table_paragraph.paragraph_format.line_spacing,
            table_profile["line_spacing_multiple"],
        )
        self.assertEqual(
            table_paragraph._p.get_or_add_pPr().find(qn("w:snapToGrid")).get(qn("w:val")),
            "0",
        )


if __name__ == "__main__":
    unittest.main()
