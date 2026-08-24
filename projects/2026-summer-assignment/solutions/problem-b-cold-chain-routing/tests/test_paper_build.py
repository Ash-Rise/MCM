"""Impact-scoped tests for the YAML-driven Problem B DOCX build."""

from __future__ import annotations

import sys
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = PROJECT_ROOT.parents[3]
sys.path.insert(0, str(REPOSITORY_ROOT / "shared"))
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from paper_format import apply_profile, load_profile  # noqa: E402
from postprocess_paper_docx import (  # noqa: E402
    TABLE_WIDTH_WEIGHTS,
    apply_project_layout,
)


PROFILE_PATH = REPOSITORY_ROOT / "shared" / "templates" / "personal-paper-profile.yaml"
DOCX_PATH = PROJECT_ROOT / "paper" / "paper.docx"


class SharedPaperFormatTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.profile = load_profile(PROFILE_PATH)

    def test_generic_formatter_reads_profile_for_page_typography_footer_and_table(self) -> None:
        document = Document()
        title = document.add_paragraph("测试题名", style="Heading 1")
        major_heading = document.add_paragraph("摘 要", style="Heading 2")
        minor_heading = document.add_paragraph("2.1 方法", style="Heading 3")
        body = document.add_paragraph("正文")
        caption = document.add_paragraph("表 1 测试表")
        reference = document.add_paragraph("[1] 测试文献")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "表头"
        table.cell(1, 0).text = "内容"

        apply_profile(document, self.profile)

        page = self.profile["page"]
        section = document.sections[0]
        self.assertAlmostEqual(section.page_width.cm, page["width_cm"], places=2)
        self.assertAlmostEqual(section.page_height.cm, page["height_cm"], places=2)
        self.assertAlmostEqual(section.left_margin.cm, page["margins_cm"]["left"], places=2)
        self.assertAlmostEqual(section.header_distance.cm, page["header_distance_cm"], places=2)
        self.assertAlmostEqual(section.footer_distance.cm, page["footer_distance_cm"], places=2)
        self.assertEqual(
            section._sectPr.find(qn("w:docGrid")).get(qn("w:linePitch")),
            str(page["document_grid_line_pitch_twips"]),
        )
        self.assertEqual(
            section.different_first_page_header_footer,
            page["different_first_page"],
        )
        self.assertIn("PAGE", section.footer._element.xml)
        self.assertIn("PAGE", section.first_page_footer._element.xml)

        typography = self.profile["typography"]
        self.assertEqual(title.runs[0].font.size, Pt(typography["title"]["size_pt"]))
        self.assertEqual(
            title.runs[0].font.color.rgb,
            RGBColor.from_string(typography["font_color_hex"]),
        )
        self.assertEqual(
            major_heading.runs[0].font.size,
            Pt(typography["heading_level_1"]["size_pt"]),
        )
        self.assertEqual(major_heading.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(
            minor_heading.runs[0].font.size,
            Pt(typography["heading_level_2_and_3"]["size_pt"]),
        )
        self.assertEqual(minor_heading.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(body.paragraph_format.line_spacing, typography["body"]["line_spacing_multiple"])
        self.assertEqual(caption.runs[0].font.size, Pt(typography["caption"]["size_pt"]))
        self.assertEqual(
            reference.runs[0].font.size,
            Pt(typography["reference_and_short_appendix"]["size_pt"]),
        )

        table_profile = self.profile["tables"]
        header_cell_properties = table.cell(0, 0)._tc.get_or_add_tcPr()
        borders = header_cell_properties.find(qn("w:tcBorders"))
        self.assertEqual(
            borders.find(qn("w:top")).get(qn("w:sz")),
            str(table_profile["top_border_ooxml_size"]),
        )
        margins = header_cell_properties.find(qn("w:tcMar"))
        self.assertEqual(
            margins.find(qn("w:start")).get(qn("w:w")),
            str(table_profile["cell_margins_twips"]["start"]),
        )
        self.assertEqual(
            table.cell(1, 0).paragraphs[0].runs[0].font.size,
            Pt(typography["table_text"]["size_pt"]),
        )

    def test_problem_b_layer_keeps_only_local_width_and_pagination_choices(self) -> None:
        document = Document()
        document.add_paragraph("测试题名", style="Heading 1")
        document.add_paragraph("一、问题重述", style="Heading 2")
        for weights in TABLE_WIDTH_WEIGHTS:
            table = document.add_table(rows=2, cols=len(weights))
            for row in table.rows:
                for cell in row.cells:
                    cell.text = "x"

        apply_profile(document, self.profile)
        apply_project_layout(document)

        first_heading = next(p for p in document.paragraphs if p.text.endswith("问题重述"))
        self.assertTrue(first_heading.paragraph_format.page_break_before)
        first_widths = [
            int(column.get(qn("w:w"))) for column in document.tables[0]._tbl.tblGrid
        ]
        self.assertEqual(len(first_widths), len(TABLE_WIDTH_WEIGHTS[0]))
        self.assertGreater(first_widths[-1], first_widths[0])
        self.assertTrue(document.tables[0].cell(0, 0).paragraphs[0].paragraph_format.keep_with_next)
        self.assertFalse(document.tables[0].cell(1, 0).paragraphs[0].paragraph_format.keep_with_next)
        self.assertFalse(document.tables[1].cell(0, 0).paragraphs[0].paragraph_format.keep_with_next)

    def test_tracked_problem_b_docx_uses_profile_without_reference_template_baggage(self) -> None:
        document = Document(DOCX_PATH)
        page = self.profile["page"]
        typography = self.profile["typography"]
        section = document.sections[0]

        self.assertAlmostEqual(section.page_width.cm, page["width_cm"], places=2)
        self.assertAlmostEqual(section.page_height.cm, page["height_cm"], places=2)
        self.assertEqual(
            section._sectPr.find(qn("w:docGrid")).get(qn("w:linePitch")),
            str(page["document_grid_line_pitch_twips"]),
        )
        self.assertEqual(
            document.paragraphs[0].runs[0].font.size,
            Pt(typography["title"]["size_pt"]),
        )
        self.assertEqual(len(document.tables), len(TABLE_WIDTH_WEIGHTS))
        self.assertEqual(len(document.inline_shapes), 3)
        with ZipFile(DOCX_PATH) as archive:
            header_parts = [
                name
                for name in archive.namelist()
                if name.startswith("word/header") and name.endswith(".xml")
            ]
            footer_parts = [
                name
                for name in archive.namelist()
                if name.startswith("word/footer") and name.endswith(".xml")
            ]
        self.assertEqual(header_parts, [])
        self.assertEqual(len(footer_parts), 2)


if __name__ == "__main__":
    unittest.main()
