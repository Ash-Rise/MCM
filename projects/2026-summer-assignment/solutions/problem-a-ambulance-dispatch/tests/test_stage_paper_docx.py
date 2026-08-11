import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = PROJECT_ROOT / "paper" / "A题论文-任务一任务二阶段稿.docx"
POSTPROCESS_PATH = PROJECT_ROOT / "src" / "postprocess_paper_docx.py"


def _load_postprocessor():
    spec = importlib.util.spec_from_file_location("postprocess_paper_docx", POSTPROCESS_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load {POSTPROCESS_PATH}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _style_num_id(style):
    p_pr = style.element.pPr
    if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
        return None
    return p_pr.numPr.numId.val


def _page_field_count(footer):
    return len(footer._element.findall(".//" + qn("w:instrText")))


def test_stage_docx_postprocessor_removes_heading_numbering_and_fixes_tables(tmp_path):
    module = _load_postprocessor()
    output = tmp_path / "stage-paper.docx"
    module.postprocess_docx(DOCX_PATH, output)

    document = Document(output)
    assert document.paragraphs[0].style.name == "Title"
    assert _style_num_id(document.styles["Heading 2"]) == 0
    assert _style_num_id(document.styles["Heading 3"]) == 0
    assert _page_field_count(document.sections[0].footer) == 1
    assert _page_field_count(document.sections[0].first_page_footer) == 1
    assert document.sections[0].footer.paragraphs[0].alignment == 1
    assert document.sections[0].first_page_footer.paragraphs[0].alignment == 1

    assert len(document.tables) == len(module.TABLE_WIDTH_WEIGHTS)
    assert document.tables[0].cell(8, 2).text == "1/h，次/h"
    for table, weights in zip(document.tables, module.TABLE_WIDTH_WEIGHTS, strict=True):
        assert table._tbl.tblPr.find(qn("w:tblStyle")) is None
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None
        table_width = int(tbl_w.get(qn("w:w")))
        grid = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid.gridCol_lst]
        assert len(grid) == len(weights)
        assert sum(grid) == table_width
        for row in table.rows:
            widths = [int(cell._tc.tcPr.tcW.get(qn("w:w"))) for cell in row.cells]
            assert widths == grid
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    assert paragraph._p.pPr.find(qn("w:pStyle")) is None
