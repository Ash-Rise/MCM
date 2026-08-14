from __future__ import annotations

import argparse
from copy import deepcopy
import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips


COMPLETE_TABLE_WIDTH_WEIGHTS = (
    # Exact dxa proportions distilled from the user-corrected v2.5 Word
    # baseline.  These values are only a fallback for newly created tables;
    # release builds use the complete table OOXML lock below.
    (1565, 5739, 1566),
    (1267, 2946, 4657),
    (2609, 1044, 1044, 1044, 1044, 1044, 1041),
    (2657, 6213),
    (1064, 1064, 1419, 1774, 1774, 1775),
    (2826, 6044),
    (2886, 2061, 2061, 2062),
    (1365, 1367, 2613, 3525),
    (1183, 1478, 2070, 4139),
    (2534, 6336),
    (1223, 1529, 1529, 2319, 2270),
    (1694, 1095, 2301, 1255, 2525),
)

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TABLE_BASELINE = (
    PROJECT_ROOT / "paper" / "v2.5" / "A题论文(v2.5).docx"
)
DEFAULT_TABLE_BASELINE_SHA256 = (
    "460f5b2953afefdbc4e506510175ddc641996158dec2006c4969ba31856fd434"
)


def table_width_weights_for_count(table_count: int) -> tuple[tuple[float, ...], ...]:
    if table_count == len(COMPLETE_TABLE_WIDTH_WEIGHTS):
        return COMPLETE_TABLE_WIDTH_WEIGHTS
    raise ValueError(f"Expected 12 tables, found {table_count}")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _validate_frozen_table_baseline(path: Path) -> None:
    if path.resolve() != DEFAULT_TABLE_BASELINE.resolve():
        return
    actual = sha256_file(path)
    if actual != DEFAULT_TABLE_BASELINE_SHA256:
        raise ValueError(
            "The frozen v2.5 table baseline changed: "
            f"expected {DEFAULT_TABLE_BASELINE_SHA256}, found {actual}. "
            "Review the Word tables and deliberately adopt a new baseline."
        )


def rebind_conversion_manifest(
    manifest_path: Path,
    output_path: Path,
    table_baseline: Path | None = None,
    allow_table_content_drift: bool = False,
) -> Path:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    source_path = Path(manifest["source"]) if manifest.get("source") else None
    if source_path is not None and source_path.exists():
        source_hash = sha256_file(source_path)
        manifest["source_sha256"] = source_hash
        manifest["project_sha256"] = source_hash
    manifest["output"] = str(output_path.resolve())
    manifest["output_sha256"] = sha256_file(output_path)
    manifest["postprocess"] = {
        "tool": "src/postprocess_paper_docx.py",
        "source_manifest": manifest_path.name,
        "table_baseline": str(table_baseline.resolve()) if table_baseline else None,
        "table_baseline_sha256": (
            sha256_file(table_baseline) if table_baseline else None
        ),
        "table_baseline_mode": (
            "hybrid_table_xml" if allow_table_content_drift else "complete_table_xml"
        ),
        "reason": (
            "Removed template numbering, preserved complete user-corrected OOXML for "
            "unchanged tables, regenerated only tables whose reviewed Markdown content "
            "changed, installed centered PAGE fields, kept the abstract on its own "
            "page, and added caption-derived alt text to body figures."
            if allow_table_content_drift
            else "Removed template numbering, preserved the user-corrected v2.5 "
            "complete table OOXML, installed centered PAGE fields, kept the abstract "
            "on its own page, and added caption-derived alt text to body figures."
        ),
    }
    output_manifest = output_path.with_suffix(".conversion.json")
    output_manifest.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return output_manifest


def _table_text(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def _copy_referenced_table_styles(target, baseline) -> None:
    style_ids = {
        style.get(qn("w:val"))
        for table in baseline.tables
        if (style := table._tbl.tblPr.find(qn("w:tblStyle"))) is not None
    }
    target_styles = target.styles.element
    baseline_styles = baseline.styles.element
    for style_id in style_ids:
        if style_id is None:
            continue
        xpath = f"{qn('w:style')}[@{qn('w:styleId')}='{style_id}']"
        existing = target_styles.find(xpath)
        if existing is not None:
            target_styles.remove(existing)
        source = baseline_styles.find(xpath)
        if source is None:
            raise ValueError(f"Missing table style {style_id!r} in baseline")
        target_styles.append(deepcopy(source))


def _replace_tables_from_baseline(target, baseline) -> None:
    """Preserve user-corrected Word table formatting when text is unchanged."""
    if len(target.tables) != len(baseline.tables):
        raise ValueError(
            f"Table count mismatch: target={len(target.tables)}, "
            f"baseline={len(baseline.tables)}"
        )
    for index, (target_table, baseline_table) in enumerate(
        zip(target.tables, baseline.tables, strict=True), start=1
    ):
        if _table_text(target_table) != _table_text(baseline_table):
            raise ValueError(f"Table {index} content differs from the formatting baseline")
    _copy_referenced_table_styles(target, baseline)
    for target_table, baseline_table in zip(
        target.tables, baseline.tables, strict=True
    ):
        target_table._tbl.getparent().replace(
            target_table._tbl, deepcopy(baseline_table._tbl)
        )


def _replace_unchanged_tables_from_baseline(target, baseline) -> list[int]:
    """Keep exact manual OOXML for unchanged tables and report changed tables."""
    if len(target.tables) != len(baseline.tables):
        raise ValueError(
            f"Table count mismatch: target={len(target.tables)}, "
            f"baseline={len(baseline.tables)}"
        )
    changed_tables: list[int] = []
    _copy_referenced_table_styles(target, baseline)
    for index, (target_table, baseline_table) in enumerate(
        zip(target.tables, baseline.tables, strict=True), start=1
    ):
        if _table_text(target_table) != _table_text(baseline_table):
            changed_tables.append(index)
            continue
        target_table._tbl.getparent().replace(
            target_table._tbl, deepcopy(baseline_table._tbl)
        )
    return changed_tables


def _replace_direct_child(parent, tag: str, source) -> None:
    current = parent.find(qn(tag))
    if current is not None:
        index = parent.index(current)
        parent.remove(current)
    else:
        index = 0
    if source is not None:
        parent.insert(index, deepcopy(source))


def _copy_paragraph_run_format(target_paragraph, baseline_paragraph) -> None:
    baseline_runs = baseline_paragraph.runs
    target_runs = target_paragraph.runs
    if not baseline_runs or not target_runs:
        return
    for index, target_run in enumerate(target_runs):
        baseline_run = baseline_runs[min(index, len(baseline_runs) - 1)]
        _replace_direct_child(target_run._r, "w:rPr", baseline_run._r.rPr)


def _copy_table_layout_from_baseline(target, baseline) -> list[int]:
    """Copy manual table layout while retaining the target table content."""
    if len(target.tables) != len(baseline.tables):
        raise ValueError(
            f"Table count mismatch: target={len(target.tables)}, "
            f"baseline={len(baseline.tables)}"
        )
    changed_tables: list[int] = []
    _copy_referenced_table_styles(target, baseline)
    for index, (target_table, baseline_table) in enumerate(
        zip(target.tables, baseline.tables, strict=True), start=1
    ):
        if len(target_table.rows) != len(baseline_table.rows):
            raise ValueError(f"Table {index} row count differs from the formatting baseline")
        if any(
            len(target_row.cells) != len(baseline_row.cells)
            for target_row, baseline_row in zip(
                target_table.rows, baseline_table.rows, strict=True
            )
        ):
            raise ValueError(f"Table {index} column count differs from the formatting baseline")
        if _table_text(target_table) != _table_text(baseline_table):
            changed_tables.append(index)

        _replace_direct_child(target_table._tbl, "w:tblPr", baseline_table._tbl.tblPr)
        _replace_direct_child(target_table._tbl, "w:tblGrid", baseline_table._tbl.tblGrid)
        for target_row, baseline_row in zip(
            target_table.rows, baseline_table.rows, strict=True
        ):
            _replace_direct_child(target_row._tr, "w:trPr", baseline_row._tr.trPr)
            for target_cell, baseline_cell in zip(
                target_row.cells, baseline_row.cells, strict=True
            ):
                _replace_direct_child(target_cell._tc, "w:tcPr", baseline_cell._tc.tcPr)
                if len(target_cell.paragraphs) != len(baseline_cell.paragraphs):
                    raise ValueError(
                        f"Table {index} paragraph structure differs from the formatting baseline"
                    )
                for target_paragraph, baseline_paragraph in zip(
                    target_cell.paragraphs,
                    baseline_cell.paragraphs,
                    strict=True,
                ):
                    _replace_direct_child(
                        target_paragraph._p,
                        "w:pPr",
                        baseline_paragraph._p.pPr,
                    )
                    _copy_paragraph_run_format(target_paragraph, baseline_paragraph)
    return changed_tables


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _ensure_child_before(parent, tag: str, *successors: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.insert_element_before(child, *successors)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def _column_widths(weights: tuple[float, ...], total_width: int) -> list[int]:
    widths = [round(total_width * weight / sum(weights)) for weight in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def _set_cell_margins(cell, *, top: int = 80, bottom: int = 80, side: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = _ensure_child_before(
        tc_pr,
        "w:tcMar",
        "w:textDirection",
        "w:tcFitText",
        "w:vAlign",
        "w:hideMark",
        "w:headers",
        "w:cellIns",
        "w:cellDel",
        "w:cellMerge",
        "w:tcPrChange",
    )
    for name, value in (("top", top), ("left", side), ("bottom", bottom), ("right", side)):
        margin = _ensure_child(tc_mar, f"w:{name}")
        margin.set(qn("w:type"), "dxa")
        margin.set(qn("w:w"), str(value))


def _set_table_geometry(table, weights: tuple[float, ...], total_width: int) -> None:
    widths = _column_widths(weights, total_width)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    table_style = tbl_pr.find(qn("w:tblStyle"))
    if table_style is not None:
        tbl_pr.remove(table_style)
    _set_width(tbl_pr, "w:tblW", total_width)
    layout = _ensure_child_before(
        tbl_pr,
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    layout.set(qn("w:type"), "fixed")
    indent = _ensure_child_before(
        tbl_pr,
        "w:tblInd",
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for column_index, cell in enumerate(row.cells):
            cell.width = Twips(widths[column_index])
            _set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_column_widths(table, weights: tuple[float, ...]) -> None:
    """Adjust column proportions without disturbing a hand-corrected table style."""
    grid = table._tbl.tblGrid
    total_width = sum(int(column.get(qn("w:w"))) for column in grid)
    widths = _column_widths(weights, total_width)
    for column, width in zip(grid, widths, strict=True):
        column.set(qn("w:w"), str(width))
    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            cell.width = Twips(widths[column_index])
            _set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[column_index])


def _set_table_borders(table) -> None:
    # Some reference templates copy table borders into every row through
    # tblPrEx.  Those row-level exceptions override a clean table-level
    # three-line definition and render as a heavy rule after every row.
    for row in table.rows:
        for exception in list(row._tr.findall(qn("w:tblPrEx"))):
            row._tr.remove(exception)
    borders = _ensure_child_before(
        table._tbl.tblPr,
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    for child in list(borders):
        borders.remove(child)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = _ensure_child(borders, f"w:{edge}")
        if edge in ("top", "bottom"):
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "10")
            border.set(qn("w:color"), "000000")
        else:
            border.set(qn("w:val"), "nil")

    header_borders = _ensure_child(table.rows[0]._tr.get_or_add_trPr(), "w:tblHeader")
    header_borders.set(qn("w:val"), "true")
    for cell in table.rows[0].cells:
        tc_borders = _ensure_child_before(
            cell._tc.get_or_add_tcPr(),
            "w:tcBorders",
            "w:shd",
            "w:noWrap",
            "w:tcMar",
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
            "w:cellIns",
            "w:cellDel",
            "w:cellMerge",
            "w:tcPrChange",
        )
        bottom = _ensure_child(tc_borders, "w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "000000")


def _remove_style_numbering(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_id = num_pr.get_or_add_numId()
    num_id.val = 0


def _format_title(document) -> None:
    paragraph = document.paragraphs[0]
    try:
        title_style = document.styles["Title"]
    except KeyError:
        title_style = document.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    paragraph.style = title_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = paragraph.text.replace("\n", "")
    paragraph.clear()
    run = paragraph.add_run(title)
    run.bold = True
    _set_run_font(run, east_asia="宋体", size=Pt(15))
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)


def _set_run_font(run, *, east_asia: str, size) -> None:
    run.font.name = "Times New Roman"
    run.font.size = size
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)


def _merge_compound_math_scripts(document) -> int:
    """Keep compound OMML subscripts and superscripts horizontal in LibreOffice."""
    merged = 0
    for script_tag in ("m:sub", "m:sup"):
        for script in document.element.body.findall(".//" + qn(script_tag)):
            runs = list(script.findall(qn("m:r")))
            if len(runs) < 2:
                continue
            texts = [run.find(qn("m:t")) for run in runs]
            if any(text is None for text in texts):
                continue
            combined = "".join(text.text or "" for text in texts)
            if re.fullmatch(r"[A-Za-z0-9=(),*+\-]+", combined) is None:
                continue
            texts[0].text = combined
            for run in runs[1:]:
                script.remove(run)
            merged += 1
    return merged


def _normalize_omml_matrix_properties(document) -> tuple[int, int]:
    """Normalize Pandoc OMML sequences that violate the ECMA child order."""
    fixed_run_properties = 0
    fixed_matrix_columns = 0
    for run_properties in document.element.body.findall(".//" + qn("m:rPr")):
        normal_text = run_properties.find(qn("m:nor"))
        script = run_properties.find(qn("m:scr"))
        style = run_properties.find(qn("m:sty"))
        if normal_text is not None and (script is not None or style is not None):
            # CT_RPR permits either `nor` or the script-style sequence, not both.
            if script is not None:
                run_properties.remove(script)
            if style is not None:
                run_properties.remove(style)
            fixed_run_properties += 1
        elif script is not None and style is not None:
            # WPS may save the script-style pair as sty,scr; ECMA requires scr,sty.
            children = list(run_properties)
            if children.index(style) < children.index(script):
                run_properties.remove(script)
                run_properties.insert(children.index(style), script)
                fixed_run_properties += 1
    for column_properties in document.element.body.findall(".//" + qn("m:mcPr")):
        count = column_properties.find(qn("m:count"))
        alignment = column_properties.find(qn("m:mcJc"))
        if (
            count is not None
            and alignment is not None
            and list(column_properties).index(count)
            > list(column_properties).index(alignment)
        ):
            # CT_MCPr requires count before mcJc.
            column_properties.remove(count)
            column_properties.insert(list(column_properties).index(alignment), count)
            fixed_matrix_columns += 1
    return fixed_run_properties, fixed_matrix_columns


def _remove_duplicate_bookmark_markers(document) -> tuple[int, int]:
    """Remove duplicate bookmark endpoints imported with locked table XML."""
    removed = []
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        seen: set[str] = set()
        count = 0
        for marker in list(document.element.body.findall(".//" + qn(tag))):
            marker_id = marker.get(qn("w:id"))
            if marker_id in seen:
                marker.getparent().remove(marker)
                count += 1
            else:
                seen.add(marker_id)
        removed.append(count)
    return removed[0], removed[1]


ARRIVAL_RATE_TERM_REPLACEMENTS = (
    ("严格4 min响应率", "4分钟内到达率"),
    ("严格4 min率", "4分钟内到达率"),
    ("严格4分钟响应率", "4分钟内到达率"),
    ("严格4分钟率", "4分钟内到达率"),
    ("严格四分钟响应率", "4分钟内到达率"),
    ("严格四分钟率", "4分钟内到达率"),
)


def _replace_text_across_runs(paragraph_element, old: str, new: str) -> int:
    """Replace visible Word text without rebuilding runs or touching OMML."""
    replaced = 0
    while True:
        text_nodes = list(paragraph_element.findall(".//" + qn("w:t")))
        values = [node.text or "" for node in text_nodes]
        combined = "".join(values)
        start = combined.find(old)
        if start < 0:
            return replaced
        end = start + len(old)
        cursor = 0
        first_index = last_index = None
        first_offset = last_offset = 0
        for index, value in enumerate(values):
            next_cursor = cursor + len(value)
            if first_index is None and start < next_cursor:
                first_index = index
                first_offset = start - cursor
            if first_index is not None and end <= next_cursor:
                last_index = index
                last_offset = end - cursor
                break
            cursor = next_cursor
        if first_index is None or last_index is None:
            raise ValueError(f"Could not map replacement span for {old!r}")
        prefix = values[first_index][:first_offset]
        if first_index == last_index:
            suffix = values[first_index][last_offset:]
            text_nodes[first_index].text = prefix + new + suffix
        else:
            suffix = values[last_index][last_offset:]
            text_nodes[first_index].text = prefix + new
            for index in range(first_index + 1, last_index):
                text_nodes[index].text = ""
            text_nodes[last_index].text = suffix
        replaced += 1


def normalize_arrival_rate_terms(document) -> int:
    replaced = 0
    for paragraph in document.element.body.findall(".//" + qn("w:p")):
        for old, new in ARRIVAL_RATE_TERM_REPLACEMENTS:
            replaced += _replace_text_across_runs(paragraph, old, new)
    return replaced


def replace_body_figure_by_caption(document, caption_prefix: str, image_path: Path) -> None:
    paragraphs = document.paragraphs
    caption_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip().startswith(caption_prefix)
        ),
        None,
    )
    if caption_index is None:
        raise ValueError(f"Could not find caption {caption_prefix!r}")
    for paragraph in reversed(paragraphs[:caption_index]):
        blip = paragraph._p.find(".//" + qn("a:blip"))
        if blip is None:
            continue
        relationship_id = blip.get(qn("r:embed"))
        document.part.rels[relationship_id].target_part._blob = image_path.read_bytes()
        return
    raise ValueError(f"Could not find image preceding caption {caption_prefix!r}")


def _disable_paragraph_grid(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    snap_to_grid = p_pr.find(qn("w:snapToGrid"))
    if snap_to_grid is not None:
        p_pr.remove(snap_to_grid)
    snap_to_grid = OxmlElement("w:snapToGrid")
    p_pr.insert_element_before(
        snap_to_grid,
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )
    snap_to_grid.set(qn("w:val"), "0")


def _set_reference_body_spacing(paragraph, multiple: float = 1.5) -> None:
    paragraph.paragraph_format.line_spacing = multiple


def _set_document_line_grid(document, line_pitch: int = 360) -> None:
    """Match the reference document's 1.5-line page grid."""
    for section in document.sections:
        section_properties = section._sectPr
        document_grid = _ensure_child(section_properties, "w:docGrid")
        document_grid.set(qn("w:type"), "lines")
        document_grid.set(qn("w:linePitch"), str(line_pitch))
        document_grid.set(qn("w:charSpace"), "0")


def _is_figure_or_table_caption(text: str) -> bool:
    return re.match(r"^[图表]\s*\d+\s+\S", text) is not None


def _format_document_typography(document) -> None:
    body_style = document.styles["Body Text"]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Title":
            continue

        if style_name in {"Heading 1", "Heading 2"}:
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(2.5)
            paragraph.paragraph_format.space_after = Pt(2.5)
            for run in paragraph.runs:
                _set_run_font(run, east_asia="黑体", size=Pt(14))
                run.bold = True
            continue

        if style_name == "Heading 3":
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(2.5)
            paragraph.paragraph_format.space_after = Pt(2.5)
            for run in paragraph.runs:
                _set_run_font(run, east_asia="宋体", size=Pt(12))
                run.bold = True
            continue

        # Pandoc may emit dangling FirstParagraph/Compact style references.
        # Normalize all non-heading content so renderers cannot inherit different line grids.
        paragraph.style = body_style

        if _is_figure_or_table_caption(text):
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            _set_reference_body_spacing(paragraph)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                _set_run_font(run, east_asia="宋体", size=Pt(10.5))
            continue

        if re.match(r"^\[\d+\]", text):
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(21)
            paragraph.paragraph_format.first_line_indent = Pt(-21)
            _set_reference_body_spacing(paragraph, 1.25)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, east_asia="宋体", size=Pt(10.5))
            continue

        if text.startswith("三问结果文件分别位于"):
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(21)
            _set_reference_body_spacing(paragraph, 1.25)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, east_asia="宋体", size=Pt(10.5))
            continue

        if paragraph._p.findall(".//" + qn("w:drawing")):
            _disable_paragraph_grid(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        if not text:
            if paragraph._p.findall(".//" + qn("m:oMath")):
                _disable_paragraph_grid(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                _set_reference_body_spacing(paragraph)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
            continue

        _disable_paragraph_grid(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_reference_body_spacing(paragraph)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        p_pr = paragraph._p.pPr
        is_list = p_pr is not None and p_pr.numPr is not None
        paragraph.paragraph_format.first_line_indent = None if is_list else Pt(24)
        if text.startswith("关键词"):
            paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            _set_run_font(run, east_asia="宋体", size=Pt(12))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _disable_paragraph_grid(paragraph)
                    _set_reference_body_spacing(paragraph, 1.15)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        _set_run_font(run, east_asia="宋体", size=Pt(10.5))


def _set_body_image_alt_text(document) -> int:
    paragraphs = document.paragraphs
    updated = 0
    for index, paragraph in enumerate(paragraphs):
        drawing_properties = paragraph._p.findall(".//" + qn("wp:docPr"))
        if not drawing_properties:
            continue
        caption = next(
            (
                candidate.text.strip()
                for candidate in paragraphs[index + 1 :]
                if candidate.text.strip()
            ),
            "",
        )
        if not caption.startswith("图"):
            continue
        for properties in drawing_properties:
            properties.set("descr", caption)
            properties.set("title", caption.split(maxsplit=1)[0])
            updated += 1
    return updated


def _keep_abstract_on_first_page(document) -> None:
    for paragraph in document.paragraphs:
        if re.match(r"^(?:1\s+|一、)问题重述", paragraph.text.strip()):
            paragraph.paragraph_format.page_break_before = True
            return
    raise ValueError("Could not find the first body heading for 问题重述")


def _set_page_field(footer) -> None:
    element = footer._element
    for child in list(element):
        element.remove(child)
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    alignment = OxmlElement("w:jc")
    alignment.set(qn("w:val"), "center")
    p_pr.append(alignment)
    paragraph.append(p_pr)
    for field_type, text in (("begin", None), (None, " PAGE "), ("separate", None)):
        run = OxmlElement("w:r")
        if field_type is not None:
            field_char = OxmlElement("w:fldChar")
            field_char.set(qn("w:fldCharType"), field_type)
            run.append(field_char)
        else:
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = text
            run.append(instruction)
        paragraph.append(run)
    value_run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = "1"
    value_run.append(value)
    paragraph.append(value_run)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph.append(end_run)
    element.append(paragraph)


def postprocess_docx(
    input_path: Path,
    output_path: Path,
    table_baseline: Path | None = DEFAULT_TABLE_BASELINE,
    allow_table_content_drift: bool = False,
) -> None:
    document = Document(input_path)
    _format_title(document)
    _merge_compound_math_scripts(document)
    _normalize_omml_matrix_properties(document)
    _set_document_line_grid(document)
    _format_document_typography(document)
    _set_body_image_alt_text(document)
    _keep_abstract_on_first_page(document)
    _remove_style_numbering(document.styles["Heading 1"])
    _remove_style_numbering(document.styles["Heading 2"])
    _remove_style_numbering(document.styles["Heading 3"])
    for section in document.sections:
        _set_page_field(section.footer)
        _set_page_field(section.first_page_footer)

    if table_baseline is not None:
        _validate_frozen_table_baseline(table_baseline)
        baseline_document = Document(table_baseline)
        if allow_table_content_drift:
            _format_generated_tables(document)
            _replace_unchanged_tables_from_baseline(document, baseline_document)
        else:
            # Exact mode is the real formatting lock: complete table elements
            # are copied last and must not be touched by later generic rules.
            _replace_tables_from_baseline(document, baseline_document)
    else:
        _format_generated_tables(document)

    _remove_duplicate_bookmark_markers(document)
    _normalize_omml_matrix_properties(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _format_generated_tables(document) -> None:
    table_width_weights = table_width_weights_for_count(len(document.tables))
    section = document.sections[0]
    content_width = int(
        round(
            (
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            / 635
        )
    )
    table_width = content_width - 200
    for table_index, (table, weights) in enumerate(
        zip(document.tables, table_width_weights, strict=True)
    ):
        current_table_width = content_width if table_index == 6 else table_width
        _set_table_geometry(table, weights, current_table_width)
        _set_table_borders(table)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                if table_index == 6:
                    _set_cell_margins(cell, top=60, bottom=60, side=35)
                    _ensure_child_before(
                        cell._tc.get_or_add_tcPr(),
                        "w:noWrap",
                        "w:tcMar",
                        "w:textDirection",
                        "w:tcFitText",
                        "w:vAlign",
                    )
                for paragraph in cell.paragraphs:
                    p_style = paragraph._p.pPr.find(qn("w:pStyle"))
                    if p_style is not None:
                        paragraph._p.pPr.remove(p_style)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    if row_index == 0:
                        paragraph.paragraph_format.keep_with_next = True
                    is_algorithm_table = table_index in {3, 5, 9}
                    if (
                        row_index != 0
                        and (
                            (is_algorithm_table and column_index in {0, 1})
                            or (table_index == 0 and column_index == 1)
                        )
                    ):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        # The user's hand-corrected v2.4 is the table-layout
                        # authority.  Apart from the symbol-description column
                        # and algorithm bodies above, every table cell is
                        # centered; do not apply the former "column 2 left"
                        # blanket rule during regeneration.
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        _set_run_font(run, east_asia="宋体", size=Pt(10))
                        if row_index == 0:
                            run.bold = True

def main() -> None:
    parser = argparse.ArgumentParser(description="Postprocess the complete paper DOCX")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--manifest", type=Path)
    parser.add_argument(
        "--table-baseline",
        type=Path,
        default=DEFAULT_TABLE_BASELINE,
        help=(
            "DOCX whose complete table XML is preserved. Defaults to the "
            "user-corrected v2.5 paper and fails if its frozen hash or any table "
            "content differs."
        ),
    )
    parser.add_argument(
        "--allow-table-content-drift",
        action="store_true",
        help=(
            "Regenerate tables whose Markdown content changed while preserving "
            "complete baseline OOXML for every unchanged table."
        ),
    )
    args = parser.parse_args()
    postprocess_docx(
        args.input,
        args.output,
        table_baseline=args.table_baseline,
        allow_table_content_drift=args.allow_table_content_drift,
    )
    if args.manifest is not None:
        rebind_conversion_manifest(
            args.manifest,
            args.output,
            table_baseline=args.table_baseline,
            allow_table_content_drift=args.allow_table_content_drift,
        )


if __name__ == "__main__":
    main()
