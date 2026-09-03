"""Mechanical acceptance checks for the current portrait DOCX candidate."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = PROJECT_ROOT.parents[1]
sys.path.insert(0, str(REPOSITORY_ROOT / "shared"))

from paper_format import validate_docx_layout  # noqa: E402


PAPER_DOCX = PROJECT_ROOT / "paper" / "paper.docx"


def _break_count(cell) -> int:
    return len(cell._tc.findall(".//" + qn("w:br")))


def test_candidate_uses_stable_shared_layout_invariants() -> None:
    document = Document(PAPER_DOCX)

    assert validate_docx_layout(document) == []
    assert len(document.sections) == 1
    assert [len(table.columns) for table in document.tables] == [3, 3, 2, 5, 5, 5, 4, 3, 3]


def test_compact_result_tables_keep_semantic_line_breaks() -> None:
    document = Document(PAPER_DOCX)

    expected = {
        3: [0, 1, 2, 1, 1],
        4: [0, 1, 2, 1, 2],
        6: [0, 2, 1, 1],
    }
    for table_index, expected_row in expected.items():
        for row in document.tables[table_index].rows[1:]:
            assert [_break_count(cell) for cell in row.cells] == expected_row
