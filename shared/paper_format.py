"""Apply and validate the repository-wide paper formatting profile."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.oxml import parse_xml, serialize_part_xml
from docx.shared import Pt, RGBColor, Twips


PROFILE_PATH = Path(__file__).resolve().parent / "templates" / "personal-paper-profile.yaml"
_ALIGNMENTS = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def load_profile(path: Path = PROFILE_PATH) -> dict[str, Any]:
    with path.open(encoding="utf-8") as stream:
        profile = yaml.safe_load(stream)
    if not isinstance(profile, dict):
        raise ValueError(f"Paper profile must be a mapping: {path}")
    return profile


def _remove_numbering(element) -> None:
    properties = element.get_or_add_pPr()
    numbering = properties.find(qn("w:numPr"))
    if numbering is not None:
        properties.remove(numbering)


def _set_run_format(run, typography: dict[str, Any], *, bold: bool | None = None) -> None:
    latin_font = typography["latin_font"]
    run.font.name = latin_font
    run.font.size = Pt(typography["size_pt"])
    run.font.color.rgb = RGBColor.from_string(typography["font_color_hex"])
    if bold is None:
        bold = typography.get("bold")
    if bold is not None:
        run.bold = bool(bold)
    italic = typography.get("italic")
    if italic is not None:
        run.italic = bool(italic)
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin_font)
    fonts.set(qn("w:hAnsi"), latin_font)
    fonts.set(qn("w:eastAsia"), typography["east_asia_font"])


def _apply_paragraph_format(paragraph, typography: dict[str, Any]) -> None:
    paragraph.alignment = _ALIGNMENTS[typography["alignment"]]
    paragraph.paragraph_format.line_spacing = typography["line_spacing_multiple"]
    paragraph.paragraph_format.space_before = Pt(typography.get("space_before_pt", 0))
    paragraph.paragraph_format.space_after = Pt(typography.get("space_after_pt", 0))
    if "snap_to_grid" in typography:
        _set_paragraph_snap_to_grid(paragraph, typography["snap_to_grid"])
    if "first_line_indent_pt" in typography:
        paragraph.paragraph_format.first_line_indent = Pt(typography["first_line_indent_pt"])
    for run in paragraph.runs:
        _set_run_format(run, typography)


def _set_paragraph_snap_to_grid(paragraph, enabled: bool) -> None:
    properties = paragraph._p.get_or_add_pPr()
    snap = properties.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        properties.insert_element_before(
            snap,
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
    snap.set(qn("w:val"), "1" if enabled else "0")


def _set_document_grid(section, grid_profile: dict[str, Any]) -> None:
    grid = section._sectPr.find(qn("w:docGrid"))
    if grid is None:
        grid = OxmlElement("w:docGrid")
        section._sectPr.append(grid)
    grid.set(qn("w:type"), str(grid_profile["type"]))
    grid.set(qn("w:linePitch"), str(grid_profile["line_pitch_twips"]))
    grid.set(qn("w:charSpace"), str(grid_profile["char_space_twips"]))


def _set_page_gutter(section) -> None:
    margins = section._sectPr.find(qn("w:pgMar"))
    if margins is None:
        raise ValueError("Section page margins were not created")
    margins.set(qn("w:gutter"), "0")


def _set_page_field(footer, typography: dict[str, Any]) -> None:
    element = footer._element
    for child in list(element):
        element.remove(child)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_formatted_run():
        run = paragraph.add_run()
        _set_run_format(run, typography)
        return run._r

    for field_type, instruction in (
        ("begin", None),
        (None, " PAGE "),
        ("separate", None),
    ):
        run = add_formatted_run()
        if field_type is not None:
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), field_type)
            run.append(field)
        else:
            text = OxmlElement("w:instrText")
            text.set(qn("xml:space"), "preserve")
            text.text = instruction
            run.append(text)
    result_run = add_formatted_run()
    result = OxmlElement("w:t")
    result.text = "1"
    result_run.append(result)
    end_run = add_formatted_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)


def _set_math_font(document, font_name: str) -> None:
    settings = document.settings._element
    math_properties = settings.find(qn("m:mathPr"))
    if math_properties is None:
        math_properties = OxmlElement("m:mathPr")
        theme_language = settings.find(qn("w:themeFontLang"))
        if theme_language is None:
            settings.append(math_properties)
        else:
            settings.insert(settings.index(theme_language), math_properties)
    math_font = math_properties.find(qn("m:mathFont"))
    if math_font is None:
        math_font = OxmlElement("m:mathFont")
        math_properties.insert(0, math_font)
    math_font.set(qn("m:val"), font_name)


def _set_math_run_font(document, font_name: str) -> None:
    """Prevent WPS from substituting Segoe Print inside native equations."""
    for math_run in document.element.findall(".//" + qn("m:r")):
        run_properties = math_run.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            insertion_index = 1 if math_run.find(qn("m:rPr")) is not None else 0
            math_run.insert(insertion_index, run_properties)
        fonts = run_properties.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), font_name)


def _find_package_part(document, partname: str):
    return next(
        (part for part in document.part.package.parts if str(part.partname) == partname),
        None,
    )


def _sync_theme_fonts(document, *, latin_font: str, east_asia_font: str) -> set[str]:
    """Replace generator-default theme fonts with profile-owned fallbacks."""
    part = _find_package_part(document, "/word/theme/theme1.xml")
    if part is None:
        return set()
    root = parse_xml(part.blob)
    replaced_fonts: set[str] = set()
    for family_name in ("majorFont", "minorFont"):
        family = root.find(".//" + qn(f"a:{family_name}"))
        if family is None:
            continue
        for element_name, font_name in (
            ("latin", latin_font),
            ("ea", east_asia_font),
            ("cs", latin_font),
        ):
            element = family.find(qn(f"a:{element_name}"))
            if element is None:
                continue
            previous = element.get("typeface")
            if previous and previous != font_name:
                replaced_fonts.add(previous)
            element.set("typeface", font_name)
    part._blob = serialize_part_xml(root)
    return replaced_fonts


def _set_explicit_rfonts(element, *, latin_font: str, east_asia_font: str) -> None:
    rfonts = element.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        element.insert(0, rfonts)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme"):
        rfonts.attrib.pop(qn(f"w:{attribute}"), None)
    for attribute, font_name in (
        ("ascii", latin_font),
        ("hAnsi", latin_font),
        ("eastAsia", east_asia_font),
        ("cs", latin_font),
    ):
        rfonts.set(qn(f"w:{attribute}"), font_name)


def _set_literal_text_color(element, color_hex: str) -> None:
    color = element.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        element.append(color)
    for attribute in ("themeColor", "themeShade", "themeTint"):
        color.attrib.pop(qn(f"w:{attribute}"), None)
    color.set(qn("w:val"), color_hex)


def _sync_style_resources(document, profile: dict[str, Any]) -> None:
    """Make inherited text styles use profile-owned fonts and literal colors."""
    typography = profile["typography"]
    common = {
        "latin_font": typography["latin_font"],
        "font_color_hex": typography["font_color_hex"],
    }
    style_profiles = {
        "Normal": dict(typography["body"], **common),
        "Body Text": dict(typography["body"], **common),
        "Title": dict(typography["title"], **common),
        "Heading 1": dict(typography["heading_level_1"], **common),
        "Heading 2": dict(typography["heading_level_2_and_3"], **common),
        "Heading 3": dict(typography["heading_level_2_and_3"], **common),
        "Heading 4": dict(typography["heading_level_2_and_3"], **common),
    }
    for style_name, style_profile in style_profiles.items():
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style_element = style.element
        rpr = style_element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style_element.append(rpr)
        _set_explicit_rfonts(
            rpr,
            latin_font=style_profile["latin_font"],
            east_asia_font=style_profile["east_asia_font"],
        )
        _set_literal_text_color(rpr, style_profile["font_color_hex"])
        style.font.size = Pt(style_profile["size_pt"])
        style.font.bold = style_profile.get("bold", False)
        style.paragraph_format.alignment = _ALIGNMENTS[style_profile["alignment"]]
        style.paragraph_format.line_spacing = style_profile["line_spacing_multiple"]
        style.paragraph_format.space_before = Pt(style_profile.get("space_before_pt", 0))
        style.paragraph_format.space_after = Pt(style_profile.get("space_after_pt", 0))
        if "first_line_indent_pt" in style_profile:
            style.paragraph_format.first_line_indent = Pt(
                style_profile["first_line_indent_pt"]
            )


def normalize_heading_hierarchy(document) -> bool:
    """Convert Pandoc's title/H2 hierarchy to the repository's Title/H1 hierarchy."""
    if not document.paragraphs:
        return False
    first = document.paragraphs[0]
    first_style = first.style.name if first.style else ""
    body_styles = {
        paragraph.style.name
        for paragraph in document.paragraphs[1:]
        if paragraph.style is not None
    }
    raw_pandoc_hierarchy = (
        first_style == "Heading 1" and "Heading 2" in body_styles
    ) or (
        first_style == "Title"
        and "Heading 1" not in body_styles
        and "Heading 2" in body_styles
    )
    if not raw_pandoc_hierarchy:
        return False
    first.style = document.styles["Title"]
    for paragraph in document.paragraphs[1:]:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Heading 2":
            paragraph.style = document.styles["Heading 1"]
        elif style_name == "Heading 3":
            paragraph.style = document.styles["Heading 2"]
        elif style_name == "Heading 4":
            paragraph.style = document.styles["Heading 3"]
    return True


def _sync_doc_defaults(document, profile: dict[str, Any]) -> None:
    """Remove theme-font inheritance from the document default run style."""
    typography = profile["typography"]
    settings = document.styles.element
    defaults = settings.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        settings.insert(0, defaults)
    rpr_default = defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        defaults.insert(0, rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    _set_explicit_rfonts(
        rpr,
        latin_font=typography["latin_font"],
        east_asia_font=typography["body"]["east_asia_font"],
    )
    _set_literal_text_color(rpr, typography["font_color_hex"])


def _collect_explicit_font_names(document) -> set[str]:
    """Collect font names that remain explicitly referenced after normalization."""
    names: set[str] = set()
    for partname in (
        "/word/document.xml",
        "/word/styles.xml",
        "/word/settings.xml",
        "/word/theme/theme1.xml",
    ):
        part = _find_package_part(document, partname)
        if part is None:
            continue
        root = parse_xml(part.blob)
        for element in root.iter():
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                value = element.get(qn(f"w:{attribute}"))
                if value:
                    names.add(value)
            if element.tag in {
                qn("m:mathFont"),
                qn("a:latin"),
                qn("a:ea"),
                qn("a:cs"),
            }:
                value = element.get("typeface") or element.get(qn("m:val"))
                if value:
                    names.add(value)
    return names


def _sync_font_table(
    document,
    *,
    removed_fonts: set[str],
    required_fonts: set[str],
    prune_unreferenced: bool = True,
) -> None:
    part = _find_package_part(document, "/word/fontTable.xml")
    if part is None:
        return
    root = parse_xml(part.blob)
    for font in list(root.findall(qn("w:font"))):
        if font.get(qn("w:name")) in removed_fonts or (
            prune_unreferenced and font.get(qn("w:name")) not in required_fonts
        ):
            root.remove(font)
    existing = {font.get(qn("w:name")) for font in root.findall(qn("w:font"))}
    for font_name in sorted(required_fonts - existing):
        font = OxmlElement("w:font")
        font.set(qn("w:name"), font_name)
        root.append(font)
    part._blob = serialize_part_xml(root)


def normalize_docx_resources(document, profile: dict[str, Any] | None = None) -> None:
    """Normalize theme, inherited styles, defaults, and the DOCX font table."""
    profile = profile or load_profile()
    typography = profile["typography"]
    replaced_theme_fonts = _sync_theme_fonts(
        document,
        latin_font=typography["latin_font"],
        east_asia_font=typography["body"]["east_asia_font"],
    )
    _sync_style_resources(document, profile)
    _sync_doc_defaults(document, profile)
    required_fonts = _collect_explicit_font_names(document)
    required_fonts.update(
        {
            typography["latin_font"],
            typography["body"]["east_asia_font"],
            typography["title"]["east_asia_font"],
            typography["heading_level_1"]["east_asia_font"],
            typography["heading_level_2_and_3"]["east_asia_font"],
            profile["equations"]["math_font"],
        }
    )
    _sync_font_table(
        document,
        removed_fonts=replaced_theme_fonts,
        required_fonts=required_fonts,
        prune_unreferenced=profile.get("docx", {}).get(
            "prune_unreferenced_font_table_entries", True
        ),
    )


def validate_docx_resources(document, profile: dict[str, Any] | None = None) -> list[str]:
    """Return resource/style violations that can trigger cross-editor substitutions."""
    profile = profile or load_profile()
    typography = profile["typography"]
    errors: list[str] = []
    theme_part = _find_package_part(document, "/word/theme/theme1.xml")
    if theme_part is not None:
        theme = parse_xml(theme_part.blob)
        expected = {
            "majorFont": (typography["latin_font"], typography["body"]["east_asia_font"], typography["latin_font"]),
            "minorFont": (typography["latin_font"], typography["body"]["east_asia_font"], typography["latin_font"]),
        }
        for family_name, expected_values in expected.items():
            family = theme.find(".//" + qn(f"a:{family_name}"))
            if family is None:
                errors.append(f"主题缺少 {family_name}")
                continue
            actual_values = []
            for element_name in ("latin", "ea", "cs"):
                element = family.find(qn(f"a:{element_name}"))
                actual_values.append(element.get("typeface", "") if element is not None else "")
            actual = tuple(actual_values)
            if actual != expected_values:
                errors.append(f"主题 {family_name} 字体不符合 profile: {actual}")

    styles_part = _find_package_part(document, "/word/styles.xml")
    if styles_part is not None:
        styles = parse_xml(styles_part.blob)
        expected_styles = {"Title", "Heading1", "Heading2", "Heading3", "Heading4"}
        for style in styles.findall(qn("w:style")):
            if style.get(qn("w:styleId")) not in expected_styles:
                continue
            color = style.find("./" + qn("w:rPr") + "/" + qn("w:color"))
            if color is None or color.get(qn("w:val")) != typography["font_color_hex"]:
                errors.append(f"样式 {style.get(qn('w:styleId'))} 未使用 profile 文字颜色")
            if color is not None and any(
                color.get(qn(f"w:{attribute}"))
                for attribute in ("themeColor", "themeShade", "themeTint")
            ):
                errors.append(f"样式 {style.get(qn('w:styleId'))} 仍继承主题颜色")

    font_table_part = _find_package_part(document, "/word/fontTable.xml")
    if font_table_part is not None:
        font_table = parse_xml(font_table_part.blob)
        explicit = _collect_explicit_font_names(document)
        explicit.update(
            {
                typography["latin_font"],
                typography["body"]["east_asia_font"],
                typography["title"]["east_asia_font"],
                typography["heading_level_1"]["east_asia_font"],
                typography["heading_level_2_and_3"]["east_asia_font"],
                profile["equations"]["math_font"],
            }
        )
        table_names = {
            font.get(qn("w:name"))
            for font in font_table.findall(qn("w:font"))
            if font.get(qn("w:name"))
        }
        stale = sorted(table_names - explicit)
        if stale:
            errors.append("fontTable.xml 含未被文档引用的字体: " + ", ".join(stale))
        if any(name.casefold().startswith("aptos") for name in table_names):
            errors.append("fontTable.xml 仍含 Aptos 字体")
    return errors


def validate_docx_layout(document, profile: dict[str, Any] | None = None) -> list[str]:
    """Return violations of stable page, heading, and body-layout invariants."""
    profile = profile or load_profile()
    page = profile["page"]
    typography = profile["typography"]
    errors: list[str] = []
    expected_section_values = {
        "page_width": page["width_twips"],
        "page_height": page["height_twips"],
        "top_margin": page["margins_twips"]["top"],
        "bottom_margin": page["margins_twips"]["bottom"],
        "left_margin": page["margins_twips"]["left"],
        "right_margin": page["margins_twips"]["right"],
    }
    for index, section in enumerate(document.sections, start=1):
        for attribute, expected in expected_section_values.items():
            actual = getattr(section, attribute).twips
            if actual != expected:
                errors.append(
                    f"第 {index} 节 {attribute}={actual}，应为 {expected} twips"
                )

    expected_alignments = {
        "Heading 1": _ALIGNMENTS[typography["heading_level_1"]["alignment"]],
        "Heading 2": _ALIGNMENTS[
            typography["heading_level_2_and_3"]["alignment"]
        ],
        "Heading 3": _ALIGNMENTS[
            typography["heading_level_2_and_3"]["alignment"]
        ],
    }
    for style_name, expected in expected_alignments.items():
        try:
            actual = document.styles[style_name].paragraph_format.alignment
        except KeyError:
            continue
        if actual != expected:
            errors.append(f"样式 {style_name} 对齐方式不符合 profile")

    in_short_appendix = False
    expected_body_spacing = typography["body"]["line_spacing_multiple"]
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            if text.startswith("附录"):
                in_short_appendix = True
            expected = expected_alignments.get(style_name)
            if expected is not None and paragraph.alignment != expected:
                errors.append(f"第 {index} 段 {style_name} 对齐方式不符合 profile")
            continue
        if (
            not text
            or index == 1
            or in_short_appendix
            or re.match(r"^[图表]\s*\d+\s", text)
            or re.match(r"^\[\d+\]", text)
            or paragraph._p.findall(".//" + qn("w:drawing"))
        ):
            continue
        actual_spacing = paragraph.paragraph_format.line_spacing
        if not isinstance(actual_spacing, (int, float)) or not abs(
            actual_spacing - expected_body_spacing
        ) < 1e-9:
            errors.append(
                f"第 {index} 段正文行距不符合 profile: {actual_spacing!r}"
            )
    return errors


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def normalize_docx_file(
    input_path: str | Path,
    output_path: str | Path | None = None,
    *,
    manifest_path: str | Path | None = None,
    overwrite: bool = False,
) -> Path:
    """Normalize one DOCX and optionally update its conversion manifest hash."""
    source = Path(input_path).resolve()
    target = Path(output_path).resolve() if output_path else source
    if not source.is_file() or source.suffix.casefold() != ".docx":
        raise FileNotFoundError(f"输入 DOCX 不存在: {source}")
    if target.exists() and target != source and not overwrite:
        raise FileExistsError(f"输出已存在: {target}")
    document = Document(str(source))
    profile = load_profile()
    apply_profile(document, profile)
    normalize_docx_resources(document, profile)
    errors = validate_docx_resources(document, profile) + validate_docx_layout(
        document, profile
    )
    if errors:
        raise ValueError("DOCX 资源规范化失败: " + "; ".join(errors))
    target.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(dir=target.parent, suffix=".docx", delete=False) as handle:
        temporary = Path(handle.name)
    try:
        document.save(str(temporary))
        os.replace(temporary, target)
    finally:
        if temporary.exists():
            temporary.unlink()
    if manifest_path is not None:
        manifest_file = Path(manifest_path).resolve()
        manifest = json.loads(manifest_file.read_text(encoding="utf-8"))
        manifest["output_sha256"] = _sha256(target)
        record = {
            "tool": "shared.paper_format.normalize_docx_resources",
            "profile": str(PROFILE_PATH),
        }
        postprocessing = manifest.setdefault("postprocessing", [])
        if record not in postprocessing:
            postprocessing.append(record)
        manifest_file.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return target


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Normalize repository paper DOCX resources")
    parser.add_argument("input", help="input DOCX")
    parser.add_argument("--output", help="output DOCX; omit to normalize in place")
    parser.add_argument("--manifest", help="optional conversion manifest to update")
    parser.add_argument("--overwrite", action="store_true", help="allow replacing a different output file")
    args = parser.parse_args(argv)
    target = normalize_docx_file(
        args.input,
        args.output,
        manifest_path=args.manifest,
        overwrite=args.overwrite,
    )
    print(target)
    return 0


def _set_cell_margins(cell, margins: dict[str, int]) -> None:
    properties = cell._tc.get_or_add_tcPr()
    container = properties.find(qn("w:tcMar"))
    if container is None:
        container = OxmlElement("w:tcMar")
        following_tags = {
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
            qn("w:headers"),
            qn("w:cellIns"),
            qn("w:cellDel"),
            qn("w:cellMerge"),
            qn("w:tcPrChange"),
        }
        insertion_index = next(
            (index for index, child in enumerate(properties) if child.tag in following_tags),
            len(properties),
        )
        properties.insert(insertion_index, container)
    for name in ("top", "start", "bottom", "end"):
        node = container.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            container.append(node)
        node.set(qn("w:w"), str(margins[name]))
        node.set(qn("w:type"), "dxa")


def _set_cell_edge(cell, name: str, *, value: str, size: int = 0) -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        following_tags = {
            qn("w:shd"),
            qn("w:noWrap"),
            qn("w:tcMar"),
            qn("w:textDirection"),
            qn("w:tcFitText"),
            qn("w:vAlign"),
            qn("w:hideMark"),
            qn("w:headers"),
            qn("w:cellIns"),
            qn("w:cellDel"),
            qn("w:cellMerge"),
            qn("w:tcPrChange"),
        }
        insertion_index = next(
            (index for index, child in enumerate(properties) if child.tag in following_tags),
            len(properties),
        )
        properties.insert(insertion_index, borders)
    edge = borders.find(qn(f"w:{name}"))
    if edge is None:
        edge = OxmlElement(f"w:{name}")
        borders.append(edge)
    edge.set(qn("w:val"), value)
    if value == "single":
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), "000000")


def _format_tables(document, profile: dict[str, Any]) -> None:
    table_profile = profile["tables"]
    text_profile = dict(profile["typography"]["table_text"])
    text_profile["latin_font"] = profile["typography"]["latin_font"]
    text_profile["font_color_hex"] = profile["typography"]["font_color_hex"]
    text_profile["snap_to_grid"] = profile["paragraphs"]["snap_to_document_grid"]
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row_index, row in enumerate(table.rows):
            row_properties = row._tr.get_or_add_trPr()
            if table_profile["prevent_row_split"] and row_properties.find(qn("w:cantSplit")) is None:
                row_properties.append(OxmlElement("w:cantSplit"))
            if (
                row_index == 0
                and table_profile["repeat_header_row"]
                and row_properties.find(qn("w:tblHeader")) is None
            ):
                row_properties.append(OxmlElement("w:tblHeader"))
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell, table_profile["cell_margins_twips"])
                for edge_name in ("top", "start", "bottom", "end", "insideH", "insideV"):
                    _set_cell_edge(cell, edge_name, value="nil")
                if row_index == 0:
                    _set_cell_edge(
                        cell,
                        "top",
                        value="single",
                        size=table_profile["top_border_ooxml_size"],
                    )
                    _set_cell_edge(
                        cell,
                        "bottom",
                        value="single",
                        size=table_profile["header_bottom_border_ooxml_size"],
                    )
                if row_index == len(table.rows) - 1:
                    _set_cell_edge(
                        cell,
                        "bottom",
                        value="single",
                        size=table_profile["bottom_border_ooxml_size"],
                    )
                for paragraph in cell.paragraphs:
                    style = paragraph._p.get_or_add_pPr().find(qn("w:pStyle"))
                    if style is not None:
                        paragraph._p.get_or_add_pPr().remove(style)
                    _remove_numbering(paragraph._p)
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    _apply_paragraph_format(paragraph, text_profile)
                    for run in paragraph.runs:
                        _set_run_format(run, text_profile, bold=(row_index == 0))


def apply_profile(document, profile: dict[str, Any] | None = None) -> None:
    """Apply generic profile rules in place; project-specific layout comes later."""
    profile = profile or load_profile()
    normalize_heading_hierarchy(document)
    page = profile["page"]
    typography = profile["typography"]
    math_font = profile["equations"]["math_font"]
    replaced_theme_fonts = _sync_theme_fonts(
        document,
        latin_font=typography["latin_font"],
        east_asia_font=typography["body"]["east_asia_font"],
    )
    _sync_font_table(
        document,
        removed_fonts=replaced_theme_fonts,
        required_fonts={
            typography["latin_font"],
            typography["body"]["east_asia_font"],
            typography["heading_level_1"]["east_asia_font"],
            math_font,
        },
    )
    _sync_style_resources(document, profile)
    _sync_doc_defaults(document, profile)
    page_number_typography = dict(
        typography["page_number"],
        latin_font=typography["latin_font"],
        font_color_hex=typography["font_color_hex"],
    )
    for section in document.sections:
        section.page_width = Twips(page["width_twips"])
        section.page_height = Twips(page["height_twips"])
        section.top_margin = Twips(page["margins_twips"]["top"])
        section.bottom_margin = Twips(page["margins_twips"]["bottom"])
        section.left_margin = Twips(page["margins_twips"]["left"])
        section.right_margin = Twips(page["margins_twips"]["right"])
        section.header_distance = Twips(page["header_distance_twips"])
        section.footer_distance = Twips(page["footer_distance_twips"])
        section.different_first_page_header_footer = page["different_first_page"]
        _set_page_gutter(section)
        _set_document_grid(section, page["document_grid"])
        _set_page_field(section.footer, page_number_typography)
        if page["different_first_page"]:
            _set_page_field(section.first_page_footer, page_number_typography)

    _set_math_font(document, math_font)
    _set_math_run_font(document, math_font)
    common_typography = {
        "latin_font": typography["latin_font"],
        "font_color_hex": typography["font_color_hex"],
        "snap_to_grid": profile["paragraphs"]["snap_to_document_grid"],
    }
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            _remove_numbering(document.styles[style_name].element)
        except KeyError:
            pass

    in_short_appendix = False
    abstract_heading_seen = False
    first_body_heading_started = False
    body_profile = dict(typography["body"], **common_typography)
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if index == 0:
            _remove_numbering(paragraph._p)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            title_profile = dict(typography["title"], **common_typography)
            _apply_paragraph_format(paragraph, title_profile)
            continue
        if style_name.startswith("Heading"):
            _remove_numbering(paragraph._p)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            key = (
                "heading_level_1"
                if style_name == "Heading 1"
                else "heading_level_2_and_3"
            )
            heading_profile = dict(typography[key], **common_typography)
            _apply_paragraph_format(paragraph, heading_profile)
            compact_text = re.sub(r"\s+", "", text)
            if compact_text == "摘要":
                abstract_heading_seen = True
            elif (
                abstract_heading_seen
                and not first_body_heading_started
                and key == "heading_level_1"
            ):
                paragraph.paragraph_format.page_break_before = bool(
                    profile["pagination"]["first_body_heading_starts_new_page"]
                )
                first_body_heading_started = True
            if text.startswith("附录"):
                in_short_appendix = True
            continue
        if paragraph._p.findall(".//" + qn("w:drawing")):
            _set_paragraph_snap_to_grid(
                paragraph,
                profile["paragraphs"]["snap_to_document_grid"],
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_with_next = True
            continue
        if re.match(r"^[图表]\s*\d+\s", text):
            caption_profile = dict(typography["caption"], **common_typography)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.keep_with_next = bool(
                re.match(r"^表\s*\d+\s", text)
            )
            _apply_paragraph_format(paragraph, caption_profile)
            continue
        if re.match(r"^\[\d+\]", text):
            reference_profile = dict(
                typography["reference_and_short_appendix"],
                **common_typography,
            )
            paragraph.paragraph_format.left_indent = Pt(reference_profile["hanging_indent_pt"])
            paragraph.paragraph_format.first_line_indent = Pt(-reference_profile["hanging_indent_pt"])
            _apply_paragraph_format(paragraph, reference_profile)
            continue
        if not text and paragraph._p.findall(".//" + qn("m:oMath")):
            _apply_paragraph_format(paragraph, body_profile)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if in_short_appendix:
            appendix_profile = dict(
                typography["reference_and_short_appendix"],
                **common_typography,
            )
            paragraph.paragraph_format.first_line_indent = Pt(0)
            _apply_paragraph_format(paragraph, appendix_profile)
            continue
        _apply_paragraph_format(paragraph, body_profile)
        if paragraph._p.get_or_add_pPr().find(qn("w:numPr")) is not None:
            paragraph.paragraph_format.first_line_indent = None
        if text.startswith("关键词"):
            paragraph.paragraph_format.first_line_indent = Pt(0)

    _format_tables(document, profile)


if __name__ == "__main__":
    raise SystemExit(main())
