"""Apply project-specific layout to the Pandoc-generated CUMCM paper DOCX."""

from __future__ import annotations

import argparse
from copy import deepcopy
import hashlib
import json
import os
import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.oxml import parse_xml


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPOSITORY_ROOT = PROJECT_ROOT.parents[1]
sys.path.insert(0, str(REPOSITORY_ROOT / "shared"))

import paper_format as shared_format  # noqa: E402


TABLE_WIDTH_WEIGHTS = (
    (0.14, 0.14, 0.72),
    (0.18, 0.62, 0.20),
    (0.42, 0.58),
    (0.03, 0.07, 0.055, 0.05, 0.05, 0.075, 0.225, 0.225, 0.22),
    (0.04, 0.075, 0.045, 0.045, 0.045, 0.05, 0.23, 0.23, 0.19, 0.05),
    (0.12, 0.20, 0.18, 0.16, 0.34),
    (0.07, 0.06, 0.055, 0.055, 0.07, 0.245, 0.245, 0.20),
    (0.08, 0.78, 0.14),
    (0.10, 0.62, 0.28),
)

LANDSCAPE_TABLE_INDICES = {3, 4, 5, 6, 7}
LANDSCAPE_RANGES = (
    ("七、问题 3", "九、问题 5"),
    ("表 6 问题 5 各无人机共享航迹与烟幕弹指派", "十、数值检验"),
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _ensure_child_before(parent, tag: str, *successors: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.insert_element_before(child, *successors)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child_before(
        parent,
        tag,
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def _set_cell_width(properties, width_dxa: int) -> None:
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.insert_element_before(
            width,
            "w:gridSpan",
            "w:hMerge",
            "w:vMerge",
            "w:tcBorders",
            "w:shd",
            "w:noWrap",
            "w:tcMar",
            "w:textDirection",
            "w:tcFitText",
            "w:vAlign",
            "w:hideMark",
            "w:headers",
            "w:cellIns",
            "w:cellDel",
            "w:cellMerge",
            "w:tcPrChange",
        )
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(width_dxa))


def _set_table_geometry(table, weights: tuple[float, ...], width_dxa: int) -> None:
    if len(table.columns) != len(weights):
        raise ValueError(f"Expected {len(weights)} columns, found {len(table.columns)}")
    widths = [round(width_dxa * weight / sum(weights)) for weight in weights]
    widths[-1] += width_dxa - sum(widths)

    table.autofit = False
    properties = table._tbl.tblPr
    _set_width(properties, "w:tblW", width_dxa)
    layout = _ensure_child_before(
        properties,
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    layout.set(qn("w:type"), "fixed")
    indent = _ensure_child_before(
        properties,
        "w:tblInd",
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    )
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell._tc.get_or_add_tcPr(), widths[index])


def _set_table_font_size(table, half_points: int) -> None:
    for run in table._tbl.findall(".//" + qn("w:r")):
        properties = run.find(qn("w:rPr"))
        if properties is None:
            properties = OxmlElement("w:rPr")
            run.insert(0, properties)
        for tag in ("w:sz", "w:szCs"):
            size = properties.find(qn(tag))
            if size is None:
                size = OxmlElement(tag)
                properties.append(size)
            size.set(qn("w:val"), str(half_points))
    for math_run in table._tbl.findall(".//" + qn("m:r")):
        properties = math_run.find(qn("w:rPr"))
        if properties is None:
            properties = OxmlElement("w:rPr")
            math_run.insert(0, properties)
        for tag in ("w:sz", "w:szCs"):
            size = properties.find(qn(tag))
            if size is None:
                size = OxmlElement(tag)
                properties.append(size)
            size.set(qn("w:val"), str(half_points))


def _set_table_cell_margins(table, horizontal_dxa: int = 45) -> None:
    properties = table._tbl.tblPr
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.insert_element_before(
            margins,
            "w:tblLook",
            "w:tblCaption",
            "w:tblDescription",
            "w:tblPrChange",
        )
    for tag, width in (
        ("w:top", 0),
        ("w:left", horizontal_dxa),
        ("w:bottom", 0),
        ("w:right", horizontal_dxa),
    ):
        margin = margins.find(qn(tag))
        if margin is None:
            margin = OxmlElement(tag)
            margins.append(margin)
        margin.set(qn("w:type"), "dxa")
        margin.set(qn("w:w"), str(width))


def _remap_heading_styles(document) -> None:
    if not document.paragraphs:
        raise ValueError("The generated document has no paragraphs")
    document.paragraphs[0].style = document.styles["Title"]
    for paragraph in document.paragraphs[1:]:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Heading 2":
            paragraph.style = document.styles["Heading 1"]
        elif style_name == "Heading 3":
            paragraph.style = document.styles["Heading 2"]


def _remove_title_rule(document) -> None:
    title = document.paragraphs[0]
    for properties in (title._p.get_or_add_pPr(), document.styles["Title"].element.get_or_add_pPr()):
        borders = properties.find(qn("w:pBdr"))
        if borders is not None:
            properties.remove(borders)


def _footnote_text_by_id(document) -> dict[str, str]:
    part = next(
        (
            related
            for related in document.part.related_parts.values()
            if str(related.partname) == "/word/footnotes.xml"
        ),
        None,
    )
    if part is None:
        return {}
    root = parse_xml(part.blob)
    return {
        footnote.get(qn("w:id")): "".join(
            node.text or "" for node in footnote.findall(".//" + qn("w:t"))
        )
        for footnote in root.findall(qn("w:footnote"))
    }


def _restore_superscript_citations(document) -> None:
    citation_text = _footnote_text_by_id(document)
    replaced = 0
    for paragraph in document.paragraphs:
        for run in list(paragraph._p.findall(qn("w:r"))):
            reference = run.find(qn("w:footnoteReference"))
            if reference is None:
                continue
            footnote_id = reference.get(qn("w:id"))
            number = citation_text.get(footnote_id, "").strip()
            if not re.fullmatch(r"\d+", number):
                raise ValueError(f"Unexpected inline footnote {footnote_id}: {number!r}")
            text = OxmlElement("w:t")
            text.text = f"[{number}]"
            run.remove(reference)
            run.append(text)
            properties = run.get_or_add_rPr()
            style = properties.find(qn("w:rStyle"))
            if style is not None:
                properties.remove(style)
            vertical = properties.find(qn("w:vertAlign"))
            if vertical is None:
                vertical = OxmlElement("w:vertAlign")
                properties.append(vertical)
            vertical.set(qn("w:val"), "superscript")

            following = run.getnext()
            if following is None or following.tag != qn("w:r"):
                raise ValueError("Citation footnote is not followed by the Markdown caret")
            following_text = following.find(qn("w:t"))
            if following_text is None or not (following_text.text or "").startswith("^"):
                raise ValueError("Citation footnote is not followed by the Markdown caret")
            following_text.text = (following_text.text or "")[1:]
            if not following_text.text:
                paragraph._p.remove(following)
            replaced += 1
    if replaced != 3:
        raise ValueError(f"Expected 3 citation markers, restored {replaced}")


def _section_properties(template, *, landscape: bool):
    properties = deepcopy(template)
    page_numbering = properties.find(qn("w:pgNumType"))
    if page_numbering is not None:
        properties.remove(page_numbering)
    section_type = properties.find(qn("w:type"))
    if section_type is not None:
        properties.remove(section_type)
    section_type = OxmlElement("w:type")
    properties.insert_element_before(
        section_type,
        "w:pgSz",
        "w:pgMar",
        "w:paperSrc",
        "w:pgBorders",
        "w:lnNumType",
        "w:pgNumType",
        "w:cols",
        "w:formProt",
        "w:vAlign",
        "w:noEndnote",
        "w:titlePg",
        "w:textDirection",
        "w:bidi",
        "w:rtlGutter",
        "w:docGrid",
        "w:printerSettings",
        "w:sectPrChange",
    )
    section_type.set(qn("w:val"), "nextPage")
    page_size = properties.find(qn("w:pgSz"))
    if page_size is None:
        raise ValueError("Section properties do not contain page size")
    if landscape:
        page_size.set(qn("w:w"), "16838")
        page_size.set(qn("w:h"), "11906")
        page_size.set(qn("w:orient"), "landscape")
    else:
        page_size.set(qn("w:w"), "11906")
        page_size.set(qn("w:h"), "16838")
        page_size.attrib.pop(qn("w:orient"), None)
    return properties


def _insert_section_break_before(paragraph, properties) -> None:
    section_paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    paragraph_properties.append(properties)
    section_paragraph.append(paragraph_properties)
    paragraph._p.addprevious(section_paragraph)


def _find_paragraph_starting(document, prefix: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def _add_landscape_result_sections(document) -> None:
    body_properties = document.element.body.sectPr
    if body_properties is None:
        raise ValueError("Document body does not contain final section properties")
    page_numbering = body_properties.find(qn("w:pgNumType"))
    if page_numbering is not None:
        body_properties.remove(page_numbering)
    portrait = _section_properties(body_properties, landscape=False)
    landscape = _section_properties(body_properties, landscape=True)
    for start_prefix, end_prefix in LANDSCAPE_RANGES:
        _insert_section_break_before(
            _find_paragraph_starting(document, start_prefix), deepcopy(portrait)
        )
        _insert_section_break_before(
            _find_paragraph_starting(document, end_prefix), deepcopy(landscape)
        )


def _set_keep_rules(document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        elif re.match(r"^表\s*\d+\s", text):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        elif re.match(r"^图\s*\d+\s", text):
            paragraph.paragraph_format.keep_together = True
        elif paragraph._p.findall(".//" + qn("m:oMath")):
            paragraph.paragraph_format.keep_together = True


def _set_figure_alt_text(document) -> None:
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        properties = paragraph._p.findall(".//" + qn("wp:docPr"))
        if not properties:
            continue
        caption = next(
            (
                candidate.text.strip()
                for candidate in paragraphs[index + 1 :]
                if candidate.text.strip()
            ),
            "",
        )
        if not caption.startswith("图"):
            raise ValueError(f"Figure at paragraph {index} is not followed by a caption")
        for item in properties:
            item.set("descr", caption)
            item.set("title", caption.split(maxsplit=1)[0])


def _request_field_update(document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def apply_layout(document) -> None:
    _remap_heading_styles(document)
    _restore_superscript_citations(document)
    shared_format.apply_profile(document)
    _remove_title_rule(document)

    profile = shared_format.load_profile()
    page = profile["page"]
    content_width = (
        page["width_twips"]
        - page["margins_twips"]["left"]
        - page["margins_twips"]["right"]
    )
    if len(document.tables) != len(TABLE_WIDTH_WEIGHTS):
        raise ValueError(
            f"Expected {len(TABLE_WIDTH_WEIGHTS)} tables, found {len(document.tables)}"
        )
    landscape_width = page["height_twips"] - 2 * page["margins_twips"]["left"]
    for index, (table, weights) in enumerate(
        zip(document.tables, TABLE_WIDTH_WEIGHTS, strict=True)
    ):
        table_width = landscape_width if index in LANDSCAPE_TABLE_INDICES else content_width
        _set_table_geometry(table, weights, table_width)
        if index in LANDSCAPE_TABLE_INDICES:
            _set_table_font_size(table, 16)
            _set_table_cell_margins(table)

    _set_keep_rules(document)
    _set_figure_alt_text(document)
    _add_landscape_result_sections(document)
    _request_field_update(document)

    title = document.paragraphs[0].text.strip()
    document.core_properties.title = title
    document.core_properties.subject = "2025 CUMCM A 题数学建模论文"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _update_manifest(manifest_path: Path, output_path: Path) -> None:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["output"] = str(output_path.resolve())
    manifest["output_sha256"] = _sha256(output_path)
    manifest.setdefault("postprocessing", []).append(
        {
            "tool": "scripts/format_paper_docx.py",
            "reason": (
                "Applied the repository formatting profile, semantic table widths, "
                "pagination keep rules, figure alt text, and anonymous metadata."
            ),
        }
    )
    output_manifest = output_path.with_suffix(".conversion.json")
    output_manifest.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--manifest", type=Path)
    args = parser.parse_args()

    document = Document(args.input)
    apply_layout(document)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        dir=args.output.parent, suffix=".docx", delete=False
    ) as temporary:
        temporary_path = Path(temporary.name)
    try:
        document.save(temporary_path)
        os.replace(temporary_path, args.output)
    finally:
        temporary_path.unlink(missing_ok=True)

    if args.manifest is not None:
        _update_manifest(args.manifest, args.output)


if __name__ == "__main__":
    main()
